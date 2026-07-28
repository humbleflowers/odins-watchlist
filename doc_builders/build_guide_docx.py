"""Generate a comprehensive, plain-English Word guide for Odin's Watchlist."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x7A, 0x54, 0x17)
INK_SOFT = RGBColor(0x55, 0x52, 0x4A)
POSITIVE = RGBColor(0x2E, 0x6B, 0x4C)
WARN = RGBColor(0x9C, 0x3E, 0x2E)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
for i, size in [(1, 20), (2, 15), (3, 12.5)]:
    h = doc.styles[f"Heading {i}"]
    h.font.size = Pt(size)
    h.font.color.rgb = RGBColor(0x1C, 0x1B, 0x17)
    h.font.bold = True


def shade(cell, hexc):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hexc)
    cell._tc.get_or_add_tcPr().append(shd)


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 5"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, htxt in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = htxt
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
        shade(c, "EFE2C6")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    doc.add_paragraph()


def bullets(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def numbered(items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def callout(title, body, kind="info"):
    color = {"info": ACCENT, "good": POSITIVE, "warn": WARN}[kind]
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22)
    r = p.add_run(title.upper())
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = color
    b = doc.add_paragraph(body)
    b.paragraph_format.left_indent = Inches(0.22)
    b.paragraph_format.space_after = Pt(12)


def why(text):
    """A 'why we did it' logic note."""
    p = doc.add_paragraph()
    r = p.add_run("Why: ")
    r.bold = True
    r.font.color.rgb = ACCENT
    p.add_run(text)


def faq(q, a):
    p = doc.add_paragraph()
    p.add_run(q).bold = True
    doc.add_paragraph(a)


# ===========================================================================
# Title
# ===========================================================================
doc.add_heading("Odin's Watchlist — The Complete Guide", level=0)
sub = doc.add_paragraph()
sr = sub.add_run("What we're doing, what we found, and the plain-English logic behind every decision")
sr.italic = True
sr.font.size = Pt(13)
doc.add_paragraph()
intro = doc.add_paragraph()
intro.add_run(
    "This guide walks through the whole project in simple terms: the tool that scores stocks, "
    "how we tested whether it actually works, the surprises we found along the way, and the "
    "concrete buy / target / stop plan it now gives you for each stock. No jargon assumed — "
    "there's a glossary at the end for any term that slips through."
)
doc.add_page_break()

# ===========================================================================
# 1. The one-paragraph summary
# ===========================================================================
doc.add_heading("1. The short version", level=1)
doc.add_paragraph(
    "You have a daily list of Indian stocks flagged by chart-pattern filters, and you wanted "
    "to find the ones most likely to jump 20%+ in the coming days-to-weeks. We built a "
    "scoring tool, then — instead of trusting it — replayed more than a year of history to "
    "measure whether its top picks actually beat random luck. They do, modestly. Then we "
    "found a much stronger signal hiding in unused price data (stocks that have quietly been "
    "beating the market), rebuilt the score around it, and that more than doubled the tool's "
    "real edge. Finally we studied the exit side — what profit target and stop-loss give the "
    "best results — and found the surprising answer that the stop matters more than the "
    "target. The tool now hands you a ranked shortlist with a specific target and stop-loss "
    "for each stock, plus an interactive dashboard with several validated LENSES (durable "
    "leadership, Wyckoff phase, catalyst gaps) for choosing between the strong names. Section "
    "11 covers the dashboard; section 12 walks through five real picks end to end; section 13 "
    "adds your RIGHTWAY Telegram group as an independent confirming vote; section 14 covers the "
    "later round of feature tests (two upgrades kept, several rejected); and section 15 covers "
    "the paper ledger that now grades the tool on its own live picks."
)
callout(
    "The single most important takeaway",
    "This tool narrows thousands of stocks down to a short daily list worth your attention. "
    "It improves your odds versus picking randomly — it does not tell you what will go up. "
    "Most individual picks still won't hit their target. Always use the stop-loss.",
    kind="good",
)

# ===========================================================================
# 2. What the daily pipeline does
# ===========================================================================
doc.add_heading("2. Where the data comes from (the daily pipeline)", level=1)
doc.add_paragraph(
    "Every trading day, an existing pipeline (run_odin.py) gathers four things and merges "
    "them into one spreadsheet named odin_DDMMYYYY.csv:"
)
bullets([
    "Chartink screeners — about 24 automated filters that flag stocks matching chart setups "
    "(breakouts, near 52-week highs, tight consolidations, and so on).",
    "NSE end-of-day prices and volumes for the whole market.",
    "“Delivery percentage” — how much of the day's trading was real buying vs. intraday churn.",
    "Bulk & block deals — large trades by big institutional players.",
])
why(
    "One stock showing up on several filters at once, with heavy real buying, is a stronger "
    "signal than any single filter alone. Merging everything into one sheet lets the scorer "
    "weigh all of it together."
)

# ===========================================================================
# 3. The scoring tool
# ===========================================================================
doc.add_heading("3. The scoring tool — what it rewards and why", level=1)
doc.add_paragraph(
    "find_swing_candidates.py reads the daily sheet and gives each stock a score from 0 to "
    "100. A higher score means more of the ingredients that, historically, came before big "
    "moves. Here's what goes into the score and the reasoning for each:"
)
table(
    ["Ingredient", "Points", "In plain terms — and why it's there"],
    [
        ["Relative strength vs. the market", "35",
         "How much the stock has out-gained the overall market (Nifty) over the last month. "
         "This is the biggest ingredient because our testing showed it was by far the most "
         "reliable predictor — a stock already beating the market tends to keep leading."],
        ["Chart patterns", "30",
         "Which screeners it triggered. Each screener's weight is now DATA-DRIVEN: we measured "
         "every screener's real hit-rate and weight it accordingly, so proven ones (52-week "
         "breakout, Darvas, range-breakout-with-volume) score high and consistently-losing ones "
         "(some COP/Fib/support screeners) score ZERO. Still capped below relative strength. "
         "See section 14."],
        ["Agreement & freshness", "15",
         "How many filters flag it at once, and whether a strong one triggered TODAY (fresh) "
         "rather than days ago (stale). Fresh, broad agreement is a stronger nudge."],
        ["Today's move & volume", "10",
         "Rewards a healthy up-day on strong volume; flags a stock already up a lot today as "
         "'chase risk' (you may be too late)."],
        ["Big-player buying", "10",
         "Rewards days where institutions bought more than they sold. Only a few stocks have "
         "this data on any given day."],
    ],
)
why(
    "The point weights aren't guesses. We measured how well each ingredient predicted real "
    "20%+ moves in history, and gave the strongest predictor (relative strength) the most "
    "weight. See section 5 for how that test worked."
)

# ===========================================================================
# 4. The big discovery: relative strength
# ===========================================================================
doc.add_heading("4. The discovery that changed everything: relative strength", level=1)
doc.add_paragraph(
    "The original scoring tool only knew which filters a stock triggered — it never looked at "
    "actual price history. While preparing the tests, we found the project already had over a "
    "year of full daily price data for nearly every stock on the market, sitting unused."
)
doc.add_paragraph(
    "We tried one simple idea on its own: rank stocks purely by how much they'd beaten the "
    "market over the past 20 trading days. That single idea predicted big moves far better "
    "than the entire original score — and, crucially, it kept working on dates it had never "
    "'seen' before (the real test that a pattern is genuine, not a fluke)."
)
table(
    ["How well it separated winners (top-10 picks/day)", "Original score", "With relative strength"],
    [
        ["On tested dates", "1.6–3.0x better than random", "3.2–9.1x better"],
        ["On NEW, unseen dates (the honest test)", "~1.1–1.4x (edge nearly gone)", "2.2–4.6x (edge holds)"],
    ],
)
callout(
    "The logic in one line",
    "“Buy what's already leading the market.” It feels uncomfortable (you're buying strength, "
    "not a bargain), and that's exactly why it works — most people hesitate to.",
    kind="good",
)
why(
    "We rebuilt the score to make relative strength the largest ingredient, then re-ran the "
    "entire year-long test to confirm the new version was genuinely better — not just "
    "different. It was: the tool's real, unseen-date edge more than doubled."
)
callout(
    "Completing the history — and the two things it proved",
    "The price archive originally started in May 2025 and had a 4-month collection hole. We "
    "backfilled 339 missing daily files straight from the exchange's own archives, so history "
    "now runs unbroken from July 2024 to today. Result one: the tool's measured edge stayed "
    "essentially the same on the completed data — strong evidence the edge was real, not an "
    "accident of missing data. Result two: every signal (relative strength, volatility-sized "
    "stops) is now available immediately for every stock, every day — no more waiting period. "
    "The deeper history also let us finally test a classic 'healthy uptrend' checklist "
    "(price above its long-term averages): it turned out largely redundant — stocks beating "
    "the market are almost always already in healthy uptrends — so it was measured, found "
    "unnecessary, and deliberately left out of the score.",
    kind="good",
)

# ===========================================================================
# 5. How we tested it (backtesting)
# ===========================================================================
doc.add_heading("5. How we checked it honestly (backtesting)", level=1)
doc.add_paragraph(
    "A score is just a theory until you check it against reality. So we replayed history: for "
    "every stock on every past day, we looked at what its price actually did next, and asked "
    "— did the stocks the tool scored highest go up more often than average?"
)
doc.add_heading("Two ways to define a 'win'", level=3)
table(
    ["Definition", "Plain meaning"],
    [
        ["Opportunity existed", "Did the price reach +20% at any point? (ignores any dip "
         "along the way — optimistic)"],
        ["Actually tradable", "Did it hit the profit target before hitting the stop-loss? "
         "(the realistic one, because you'd have a stop in place)"],
    ],
)
doc.add_heading("The honesty checks we built in", level=3)
bullets([
    "New-date test: we split history into an earlier part and a later part, and checked the "
    "edge held up on the later dates the tool was never tuned on. (Anything can look great on "
    "the dates you fitted it to — the later dates are the truth.)",
    "Gap-awareness: there was a 4-month hole in the data collection, and the tool was built "
    "to never treat a price from before a gap as if it came right after it. (The hole has "
    "since been backfilled entirely from the exchange's archives — see section 4 — but the "
    "guard stays, in case collection ever lapses again.)",
    "Corporate-action guard: giant one-day jumps (from stock splits/bonus issues, not real "
    "moves) are thrown out.",
    "Only realistic stocks: we ignore illiquid penny stocks and ETFs/index funds — you "
    "couldn't reliably trade those anyway.",
])
why(
    "It's easy to build a tool that looks brilliant on past data and fails on real money. "
    "Every check above exists to stop us fooling ourselves. When results were weaker than an "
    "early buggy run suggested, we reported the weaker honest number."
)

# ===========================================================================
# 6. Problems we found and fixed
# ===========================================================================
doc.add_heading("6. Problems we found and fixed along the way", level=1)
callout(
    "Bug 1 — Index funds pretending to be stocks (fixed)",
    "The sector list labelled 176 ETFs/index funds (NIFTYBEES, BANKBEES, etc.) as ordinary "
    "'Equity' stocks. They were quietly inflating the results with impossible-looking numbers. "
    "Now filtered out everywhere.",
    kind="good",
)
callout(
    "Bug 2 — Delivery data silently vanishing (fixed)",
    "A behind-the-scenes date-matching mismatch meant the 'delivery percentage' data was "
    "silently missing from every recent daily sheet — no error, just blank. Traced to a "
    "timezone quirk and fixed; the data now flows through again.",
    kind="good",
)
callout(
    "Bug 3 — ~300 tradeable stocks quietly ignored (fixed)",
    "The price engine was only reading stocks in the 'EQ' (normal) settlement series. But the "
    "exchange also lists regular equities under 'BE' (delivery-only settlement) and 'BZ' "
    "(delivery-only + surveillance) — and about 300 of the watchlist's stocks live there, "
    "including names that MOVED from EQ to BE (like DEEDEV). These were silently dropped, so "
    "they got no relative-strength or volatility signals — and one, DEEDEV, was briefly ranked "
    "#1 off months-old leftover data and even compared against a date BEFORE its own signal. "
    "Fixed two ways: include EQ + BE + BZ stocks, and skip any stock whose data is genuinely "
    "too old (for real delistings).",
    kind="good",
)
callout(
    "A trade-off that came with Bug 3's fix",
    "Restoring those ~300 stocks raised the tool's measured edge — but many BE/BZ names are "
    "small, jumpy, delivery-only stocks that spike (and drop) more easily. About 15% of the "
    "daily top picks are now these Trade-to-Trade names, and they account for an outsized "
    "share of the 'wins' in testing. They're legitimate and tradeable, but carry extra risk "
    "(surveillance, price-band limits, thinner liquidity), so the headline numbers are a bit "
    "rosier than what you'd realistically capture. The mainstream 'EQ' stocks (~82% of picks) "
    "remain the reliable core; scrutinise BE/BZ names extra carefully on the chart.",
    kind="warn",
)

# ===========================================================================
# 7. The exit study: target & stop
# ===========================================================================
doc.add_heading("7. The exit question: what target and stop are best?", level=1)
doc.add_paragraph(
    "The tool was built around a +20% target. But is that the best choice? We tested a whole "
    "range of profit targets, each paired with a stop-loss, on the tool's daily top-10 picks. "
    "The right way to judge 'best' is expected profit per trade — because a small target gets "
    "hit more often but earns little, so a high hit-rate alone is misleading."
)
doc.add_heading("What we found — the target is a weak lever", level=3)
table(
    ["Profit target (with a −8% stop)", "How often hit", "Average profit per trade"],
    [
        ["+5%", "49% (sounds great!)", "−0.99% (actually the worst — risking 8 to make 5)"],
        ["+20% (the original)", "20%", "+0.26%"],
        ["+30%", "13%", "+0.80%"],
        ["+40% (best target alone)", "8%", "+1.03%"],
    ],
)
doc.add_paragraph(
    "Two things stand out. First, small targets LOSE money — the tempting +5% target is the "
    "worst earner, because risking 8% to make 5% is a losing trade even though it 'wins' "
    "often. Second, once the target is wide enough (15%+), pushing it wider barely moves the "
    "needle — they all cluster in a narrow band. The target just isn't the main lever."
)
doc.add_heading("The real lever is the stop-loss", level=3)
doc.add_paragraph("Holding the target at +30% and simply tightening the stop roughly doubles the profit:")
table(
    ["Setup", "Average profit per trade", "Profit factor"],
    [
        ["+30% target / −8% stop", "+0.80%", "1.17"],
        ["+30% target / −6% stop", "+1.34%", "1.34"],
        ["+30% target / −5% stop / exit within ~2 weeks", "+1.49% (best)", "1.52"],
    ],
)
callout(
    "The logic — why cutting losses fast wins here",
    "These picks behave like breakouts: lots of small losers, a few big winners. That payoff "
    "rewards cutting losers quickly and letting winners run to a wide target — NOT a modest, "
    "even target. So: tight stop, wide target, don't hold a stalled trade forever.",
    kind="good",
)
why(
    "The exact profit-per-trade numbers shifted a little when we restored those ~300 extra "
    "stocks (they include volatile names with big winners), but the lesson was identical "
    "before and after — which is exactly why we trust it. Across every version of the test, "
    "the stop-loss and holding period mattered far more than the target. The direction is the "
    "robust finding; treat the precise percentages as approximate."
)

# ===========================================================================
# 8. What the tool now gives you per stock
# ===========================================================================
doc.add_heading("8. What the tool now gives you for each stock", level=1)
doc.add_paragraph("Putting it all together, each candidate now comes with a concrete plan:")
bullets([
    "A Setup Score (0–100) and a plain-English reason it scored that way.",
    "A suggested Target and Stop-loss price, sized to that specific stock's volatility.",
    "For older sheets: what actually happened since — the current price, the % move since the "
    "signal, and the highest price it reached (so you can see how past picks played out).",
])
doc.add_heading("How the per-stock target & stop are set", level=3)
doc.add_paragraph(
    "A calm, low-volatility stock and a jumpy, high-volatility one shouldn't use the same "
    "5% stop — 5% might be a single day's normal wiggle for the jumpy one. So the stop is "
    "sized to each stock's own daily volatility (its 'ATR'), then kept within sensible limits, "
    "and the target is set at 5x the risk:"
)
table(
    ["Stock's daily swing (ATR)", "Stop-loss", "Target", "Reward : Risk"],
    [
        ["Calm (≤3.5%)", "−5%", "+25%", "5 : 1"],
        ["Medium (~5%)", "−7.5%", "+38%", "5 : 1"],
        ["Jumpy (~8%)", "−10%", "+40%", "4 : 1"],
        ["Data too fresh to measure", "−6%", "+30%", "5 : 1 (fallback)"],
    ],
)
why(
    "The typical stock lands on a −5% stop / +25% target — exactly the sweet spot the exit "
    "study found. Volatility only widens the stop for genuinely jumpy names so they aren't "
    "stopped out on normal noise. These are a sensible starting framework, not a guarantee."
)
doc.add_heading("Example of the output", level=3)
ex = doc.add_paragraph()
for line in [
    "#2  KALYANKJIL  Kalyan Jewellers   Score:70   Breakout / Momentum   Price:476.15",
    "     Plan: entry Rs476.15 | target Rs619.00 (+30%) | stop Rs447.58 (-6%) | R:R 1:5.0",
    "     Since signal (2026-07-10): Rs476.15 -> Rs570.00 (+19.7%)",
    "     Peak since signal: Rs574.40 on 2026-07-17 (+20.6% from signal)",
]:
    r = ex.add_run(line + "\n")
    r.font.name = "Consolas"
    r.font.size = Pt(9)

# ===========================================================================
# 9. What to actually do
# ===========================================================================
doc.add_heading("9. How to actually use it (and what not to expect)", level=1)
numbered([
    "Run 'python run_odin.py' once after the close — it now does the whole chain automatically: "
    "builds the sheet, pulls new RIGHTWAY Telegram messages, scores the shortlist, updates the "
    "paper ledger, and rebuilds the dashboard. You get a ranked shortlist with a plan per stock.",
    "Treat the shortlist as candidates to chart and judge yourself, not buy signals.",
    "If you take a trade, set the stop-loss BEFORE you buy, and respect it. This is the single "
    "biggest driver of whether the edge turns into profit.",
    "Expect most picks not to reach target. The tool improves your odds; it doesn't remove risk.",
    "Size positions so a string of small stop-outs is survivable — that's the normal pattern here.",
])
callout(
    "Honest limitations",
    "Results use daily closing prices (not every intraday spike), ignore brokerage/slippage "
    "(real returns a bit lower), and the watchlist test covers one ~13-month window — one "
    "market environment, not many. (The underlying price history is now unbroken from July "
    "2024 to today, so all signals are available immediately; that earlier 'needs a month of "
    "recent data' limitation is gone.) Roughly 15% of the top picks are small, volatile "
    "delivery-only (BE/BZ) stocks whose backtested numbers look better than you'd "
    "realistically capture. Trust the DIRECTION of these findings — and the mainstream 'EQ' "
    "core — more than any single exact number.",
    kind="warn",
)

# ===========================================================================
# 10. The pieces (files)
# ===========================================================================
doc.add_heading("10. The pieces, in order you'd use them", level=1)
table(
    ["File", "What it's for"],
    [
        ["run_odin.py", "The daily pipeline. Builds today's raw watchlist AND then auto-runs "
         "the whole downstream chain: Telegram delta-fetch, scorer, paper-ledger update, "
         "dashboard rebuild, ledger snapshot. Run this one command and everything else follows "
         "(use --no-downstream to build only the sheet)."],
        ["find_swing_candidates.py", "Scores the watchlist and prints your ranked shortlist "
         "with per-stock target & stop. This is the one you run daily. It also pulls new "
         "RIGHTWAY Telegram messages first (best-effort)."],
        ["config/screener_weights.csv", "The DATA-DRIVEN screener weights the scorer uses "
         "(section 14). Regenerate from the backtest with feature_backtests/gen_screener_weights.py."],
        ["technical_indicators.py", "Computes the real price-based signals (relative strength, "
         "volatility) from raw market history. Feeds the scorer."],
        ["backtest_swing_candidates.py", "Replays history to measure whether the tool's edge "
         "is real. Run occasionally to re-check."],
        ["optimize_target.py", "The exit study — finds the best target/stop combination."],
        ["backfill_bhavcopy.py", "One-shot repair tool: downloads any missing daily price "
         "files from the exchange archives (used to fill the 4-month gap and extend history "
         "back to July 2024). Rerunnable any time collection lapses."],
        ["make_dashboard.py", "Turns the daily shortlist into the interactive triage "
         "dashboard (dashboard.html) — tiers, lenses, leadership strips, filters, help."],
        ["telegram_fetch.py", "End-of-day reader for your RIGHTWAY Telegram group and its "
         "sub-groups. Uses your own account (Telethon) to pull new messages into "
         "telegram_messages.csv — works even when the group blocks the Export button. Run "
         "after market close. See section 13."],
        ["parse_telegram_recs.py", "Turns those raw Telegram messages into structured picks "
         "(symbol, direction, entry, target, stop) in telegram_recs.csv, using the NSE symbol "
         "list to pull tickers out of free text."],
        ["fetch_telegram_delta.py", "Safe wrapper that pulls only NEW Telegram messages; the "
         "scorer runs it automatically each day. Never breaks the run if Telegram is offline "
         "or unconfigured."],
        ["paper_ledger.py", "The paper-trading ledger — snapshots each day's shortlist and "
         "tracks how the tool's OWN picks actually perform, per lens. See section 15."],
    ],
)

# ===========================================================================
# 11. The dashboard & its lenses
# ===========================================================================
doc.add_heading("11. The daily dashboard and its lenses", level=1)
doc.add_paragraph(
    "make_dashboard.py builds an interactive page (dashboard.html) from the day's shortlist. "
    "Beyond the score, it adds several LENSES — each one is a way of looking at the same "
    "stocks, and each was put through the same backtest before being trusted. The golden "
    "rule that emerged: relative strength already ranks the strong names to the top, so these "
    "extra signals are NOT added to the score — they're how you choose BETWEEN the strong "
    "names."
)
bullets([
    "Conviction tiers (A / B / C) — a six-check summary of how textbook each setup is.",
    "The 20-day leadership strip — has this stock STAYED a market leader, or just flared up "
    "once? Validated: persistent leaders (led 12+ of last 20 days) hit target ~24% vs ~15% "
    "for fresh ones.",
    "Wyckoff lens (SOS / LPS / Base) — where the stock is in the accumulation-to-markup "
    "cycle, read from which screeners it fired.",
    "Episodic Pivot (EP) — a gap-up on heavy volume (a catalyst). Validated: ~2.5x the base "
    "rate standalone, ~24% vs ~17% within the top names. A different axis from strength — "
    "'did something happen today', not 'is it already strong'. REFINED (see section 14): the "
    "EP badge now fires only when the gap HELD — closed in the top half of its day's range — "
    "which backtested at ~19% vs ~12% for gaps that faded. A held gap means real buyers showed "
    "up; a faded gap is no longer flagged.",
    "Emerging (EMRG) — a volume THRUST (>=3x the stock's own normal) while relative strength "
    "is rising but still modest (0 to +25). It tries to catch a quiet stock the moment it "
    "wakes up, BEFORE it's an obvious leader (the STALLION-type early move). Validated at "
    "~1.5x the base rate — real, but roughly HALF the edge of a confirmed leader (~3x), so "
    "treat it as an earlier, lower-confidence entry, not a better one.",
    "Relative-volume readout — every stock now shows how many times its own normal volume it "
    "traded today (3x = three times average), so you can tell genuine participation from a "
    "quiet drift at a glance. It also powers the EP and Emerging lenses.",
    "Cohort filters — Persistent leaders, New leaders today, Emerging, and the setup lenses — "
    "each a one-click shortlist.",
])
callout(
    "What we tested and deliberately REJECTED (the method working)",
    "Not everything passed. The Minervini trend template and NSE delivery-% accumulation were "
    "both built, backtested, and dropped — they didn't beat relative strength and, in "
    "delivery's case, the high-delivery names actually continued slightly WORSE. Both are "
    "'quality/strength' proxies that just re-measure what relative strength already captures. "
    "The lenses that survived (persistence, Wyckoff breakout, episodic pivot) all measure "
    "something DIFFERENT — durability, structure, or a catalyst. That's the pattern: add a "
    "new axis, not another strength proxy.",
    kind="warn",
)
doc.add_heading("Working the dashboard — the interactive tools", level=3)
doc.add_paragraph(
    "Beyond the lenses, the page is now a full triage workspace. Everything below lives in the "
    "one self-contained dashboard.html — no server, and your settings, marks and notes are "
    "remembered in the browser between days."
)
bullets([
    "Four tabbed views — SHORTLIST (the main table), WATCHLIST (names you've starred), CHANGES "
    "(what's new / dropped / moved since yesterday's sheet), and LEDGER (your open paper "
    "positions). Switch with the tabs or the number keys 1–4.",
    "Mark a pick — click a row's star once for 'watching', twice for 'taken'; shift-click to "
    "dismiss a name you've ruled out. Marks persist and populate the Watchlist tab.",
    "One-click chart — the 'TV' link on every symbol opens it on TradingView.",
    "Position sizer — type your capital and risk-% once (top bar) and every row shows how many "
    "shares to buy so a stop-out costs exactly that risk. Sizes update live.",
    "A price sparkline and the stock's sector on every row; a NEW badge and a rank-move arrow "
    "versus yesterday; and, if a name is in your paper ledger, an inline 'held +x%' / "
    "'target' / 'stopped' badge.",
    "Score composition — expand a row to see an indicative bar of what drove its score "
    "(relative strength vs. patterns vs. freshness vs. volume).",
    "Market-context strip and alerts — the top of the page shows market breadth (share of the "
    "whole market above its 50-day average, CONTEXT only — a breadth filter tested worse) and "
    "flags any ledger position that just hit target/stop or is near one.",
    "Convenience — light/dark theme toggle, compact-density mode, show/hide columns, a sector "
    "filter, keyboard navigation (/ to search, j/k to move, Enter to expand), and Export-CSV "
    "of whatever you're currently viewing.",
])
callout(
    "How the paper ledger and the dashboard connect",
    "Each day the pipeline snapshots the shortlist into the paper ledger (section 15) and "
    "resolves prior picks against the exchange's prices. The dashboard then reads that ledger "
    "back — the P&L panel, the Ledger tab, the per-row 'held' badges and the near-target/stop "
    "alerts all come from it. So the tool visibly grades itself on its own picks, right where "
    "you triage.",
    kind="good",
)

# ===========================================================================
# 12. How to select — worked examples
# ===========================================================================
doc.add_heading("12. How to select — five worked examples", level=1)
doc.add_paragraph(
    "The recipe in one paragraph: start from the daily top ~10 by score (the edge lives in "
    "the top slice, not in score decimals). Prefer strong relative strength (≥ +15), a fresh "
    "breakout today, volume confirmation, and — the tiebreakers — durable leadership (a full "
    "strip) or a fresh catalyst (EP). Drop anything already up &gt; 15% today (you're "
    "chasing) or with no leadership history (dead money). Then open the chart yourself and "
    "confirm the entry. Here are five real names from the 22 Jul 2026 sheet, each showing a "
    "different lesson:"
)
doc.add_paragraph()

def example(name, verdict, vcolor, lines):
    p = doc.add_paragraph()
    r = p.add_run(name + "  ")
    r.bold = True; r.font.size = Pt(12)
    v = p.add_run(verdict)
    v.bold = True; v.font.color.rgb = vcolor; v.font.size = Pt(10.5)
    for lab, txt in lines:
        q = doc.add_paragraph()
        q.paragraph_format.left_indent = Inches(0.25)
        rr = q.add_run(lab + ": "); rr.bold = True
        q.add_run(txt)
    doc.add_paragraph()

example("GANDHAR — the clean ideal", "ACCEPT", POSITIVE, [
    ("What the tool shows", "Tier A, 6/6 checks (the only 6/6 today). Score 91, RS +31pp, "
     "fired a fresh DARVAS + 52-week + range breakout today, up a healthy +4.0%, led on 11 "
     "of the last 20 days."),
    ("Why it's the textbook buy", "Every box ticks WITHOUT any red flag — strong-but-not-"
     "extreme move (not a chase), a genuine fresh breakout, decent leadership behind it."),
    ("The plan", "Entry ~₹236 → target ₹317 (+34%), stop ₹220 (−6.8%), reward:risk ~1:5."),
    ("Your job", "Confirm on the chart it's breaking a real base, then take the printed stop."),
])
example("HUHTAMAKI — the premium leader", "ACCEPT (manage entry)", POSITIVE, [
    ("What the tool shows", "Tier A, RS +46pp, and a SOLID 20/20 leadership strip — a "
     "persistent leader — firing a fresh 52-week breakout today. Phase: SOS."),
    ("Why it's premium", "Persistent leadership + a fresh breakout is the highest-conviction "
     "combination the tool can show (the ~24% cohort)."),
    ("The one caution", "It's already up +10.5% today on 8x volume — strong, but you're not "
     "early. Consider a smaller size or waiting for a tight pullback rather than chasing the "
     "green candle. Plan: entry ~₹292 → target ₹391 (+34%), stop ₹272 (−6.8%)."),
])
example("EMAMIPAP — high score, still a pass", "REJECT / high caution", WARN, [
    ("What the tool shows", "Ranks #4 with a big score (92) and RS +38 — looks great at a "
     "glance."),
    ("Why it's actually a trap", "Two red flags the rank hides: its leadership strip is "
     "0/20 (it has NEVER been a market leader in 20 days — no durable trend, this is a "
     "one-day flare), AND it's already up +20% TODAY (the 'ext' chase-risk flag). You'd be "
     "buying a spike with no history behind it."),
    ("The lesson", "A high score is necessary, not sufficient. The leadership strip and the "
     "chase flag are exactly what save you here — always glance at them before the price."),
])
example("SUPREMEINF — the catalyst play", "SPECULATIVE / watch", RGBColor(0x7A, 0x54, 0x17), [
    ("What the tool shows", "The day's only Episodic Pivot — gapped +4.4% on 5.8x volume "
     "(EP↑ badge), RS +22. But a 0/20 leadership strip and a lower score (66, rank 48)."),
    ("How to read it", "A fresh catalyst with no leadership history — the OPPOSITE profile "
     "to HUHTAMAKI. The gap says 'something happened today'; the empty strip says 'unproven'."),
    ("If you take it", "Treat it as speculative: smaller size, and only if the chart shows "
     "the gap came off a tight base and is holding. Catalyst moves can continue hard or fade "
     "fast — the stop matters more than usual."),
])
example("TIRUPATIFL — a leader resting", "WATCHLIST (not today)", RGBColor(0x7A, 0x54, 0x17), [
    ("What the tool shows", "The strongest RS on the board (+67pp) and a persistent 17/20 "
     "leadership strip — but phase is BASE, not SOS: it's consolidating, not breaking out "
     "today."),
    ("Why wait", "A proven leader catching its breath. There's no fresh trigger to act on "
     "right now, so buying today is guessing. Add it to your watchlist."),
    ("What to watch for", "The day its strip stays lit AND a fresh SOS breakout fires — that "
     "'long base → ignition' is the setup you want to catch, not the mid-base drift."),
])

doc.add_paragraph(
    "The thread through all five: the SCORE gets a name onto the list; the LENSES (leadership "
    "strip, phase, EP, chase flags) tell you which of the high scorers is actually a buy, a "
    "wait, or a trap — and the CHART is always the final confirmation before you risk money."
)

# ===========================================================================
# 13. Telegram / RIGHTWAY confluence
# ===========================================================================
doc.add_heading("13. Bringing in your Telegram group (RIGHTWAY confluence)", level=1)
doc.add_paragraph(
    "You're a member of the RIGHTWAY trading community on Telegram, where an admin posts calls "
    "and members discuss stocks. We wired that chatter into the dashboard as one more, "
    "INDEPENDENT vote — never a buy signal on its own. The idea is CONFLUENCE: when the tool "
    "flags a strong setup AND your group is already talking about the same name, that's a "
    "second, unrelated opinion pointing the same way."
)
doc.add_heading("How the pipeline reads Telegram", level=3)
numbered([
    "telegram_fetch.py logs in once with your own Telegram account (credentials come from "
    "environment variables, never saved to a file) and reads every group whose title starts "
    "with RIGHTWAY, including its sub-groups and forum topics. It works even though the group "
    "blocks the app's Export button — that only stops the button, not a client reading "
    "messages your account can already see. Image tips still contribute their caption text "
    "(where the stock name usually is). New messages are appended to telegram_messages.csv, "
    "and the last-seen message id per group is remembered (telegram_state.json) so re-runs "
    "only pull what's new. Run it once at end of day.",
    "parse_telegram_recs.py reads that message log and pulls out structured picks — symbol, "
    "direction, entry, target, stop-loss — by matching tokens against the NSE symbol list "
    "(telegram_recs.csv).",
    "make_dashboard.py cross-references those mentions against the day's shortlist and shows "
    "them in the RW column and two filters.",
])
doc.add_heading("What you see on the dashboard", level=3)
bullets([
    "RW column — how many times the group mentioned this stock in the last ~45 days (blank = "
    "not in recent buzz).",
    "A “^” marker — the ADMIN (the broadcast channel that posts the actual calls) has called "
    "it at some point; hover for the date. Expand the row to see the admin's stated target.",
    "A bold amber “*” — the admin flagged it important recently (a POSITIONAL / MULTIBAGGER / "
    "LONG TERM / FOCUS conviction tag within the last ~4 months).",
    "Two filters — “In RIGHTWAY” (mentioned by anyone) and the tighter, higher-value “Admin "
    "call” (every stock the admin channel actually called, with a target or conviction tag).",
])
callout(
    "Read this channel honestly — it is promotional",
    "This group is heavily promotional and often posts winners AFTER they've already moved, so "
    "a mention means “they're discussing it,” NOT that it's a good buy. About 43 of 60 "
    "dashboard names get mentioned on a given day, so a plain mention isn't selective — the "
    "admin “^” / “*” markers are far more meaningful than raw count. Most admin calls in this "
    "data are 9–24 months old, so treat their targets as historical context, not live levels; "
    "a recent call shows the live upside, an old one just shows the target and the date it was "
    "made. The admin's target is the ADMIN'S stated goal — the tool's own volatility-based "
    "target and stop are still the only ones with a backtest behind them.",
    kind="warn",
)
why(
    "Use it strictly as a confirming nudge on a stock the tool already rates — especially "
    "‘strong setup + fresh ADMIN call’ — never as a reason to buy something the tool doesn't "
    "like. We have SINCE backtested whether Telegram mentions actually precede moves (section "
    "14): they beat the base rate only slightly (~12% vs ~9%), confirming this is confluence, "
    "NOT a validated edge. Keep treating it as a nudge, never a signal."
)

# ===========================================================================
# 14. Feature backtests — what we tested, kept, and rejected
# ===========================================================================
doc.add_page_break()
doc.add_heading("14. What we tested next — and the two upgrades that survived", level=1)
doc.add_paragraph(
    "After the tool was working, we drew up a list of candidate features and put every one "
    "through the same honest test: does it separate winners AMONG the already-strong names "
    "(the daily shortlist), rather than just re-sorting by strength? Most ideas failed that bar "
    "— and knowing what does NOT work is as valuable as knowing what does. Two upgrades passed "
    "and are now built in."
)
doc.add_heading("The two that were built", level=3)
table(
    ["Upgrade", "What changed", "Evidence"],
    [
        ["Data-driven screener weights", "The score used to weight chart screeners by hand. "
         "Now each screener's weight comes from its MEASURED hit-rate, and consistently-losing "
         "screeners are set to zero (dropped).", "The per-screener edge held up out-of-sample "
         "(the good/bad ranking repeated in a later period, rank correlation 0.75) — so this is "
         "a real, stable signal, not curve-fitting."],
        ["Episodic-Pivot follow-through", "An EP gap now only earns its badge if it HELD — "
         "closed in the top half of the day's range. A gap that faded to the lows is no longer "
         "flagged.", "Held gaps hit target ~19% vs ~12% for faded ones — a +7-point edge from a "
         "one-line rule."],
    ],
)
doc.add_heading("The ones we tested and REJECTED (the method working)", level=3)
bullets([
    "Market-regime / breadth filter — a 'only trade when the market is strong' rule would "
    "actually have HURT (entries on weak-breadth days did better). Rejected.",
    "Sector rotation (buy leaders in leading sectors) — only a ~2-point edge. Too small to add.",
    "Delivery-percentage TREND — no edge (~-3 pts), confirming delivery adds nothing beyond the "
    "strength the score already captures.",
    "Volatility 'squeeze' / coil — no edge (~-2 pts) in this data.",
    "Following bulk/block 'smart-money' buys — no edge (~-6 pts); those trades are often "
    "reported AFTER the move, and most big 'buyers' are just high-frequency desks.",
    "Trailing / scale-out exits — our current fixed target-and-stop is already near-best; "
    "trailing stops cut the winners short and LOWERED profit. Kept the simple exit.",
    "Telegram mentions as a scored signal — only ~12% vs ~9% base. Kept as confluence only "
    "(section 13), never folded into the score.",
])
callout(
    "The pattern behind every accept and reject",
    "Everything that PASSED (persistence, episodic pivot, the screener re-weighting) measures a "
    "DIFFERENT axis — durability, a catalyst, or which specific pattern actually works. "
    "Everything REJECTED was either another proxy for strength the score already has, or a "
    "signal that simply didn't predict. When judging any future idea, ask: does it add a new "
    "axis, or re-measure strength? Only the former helps. Full scripts and numbers live in "
    "feature_backtests/.",
    kind="good",
)

# ===========================================================================
# 15. The paper ledger — measuring the tool on itself
# ===========================================================================
doc.add_heading("15. The paper ledger — grading the tool on its own picks", level=1)
doc.add_paragraph(
    "The backtest tells us how the method did on HISTORY. The paper ledger tells us how it's "
    "doing on the picks it makes GOING FORWARD — live, honestly, and with no cherry-picking. "
    "Every day the pipeline snapshots the shortlist into a running book (paper_ledger.py), then "
    "checks each open pick against the official exchange prices and marks the outcome."
)
bullets([
    "Each pick is recorded with its entry, volatility-sized target and stop, and its lens tags "
    "(persistent leader / episodic pivot / emerging / fresh breakout / RIGHTWAY admin).",
    "Outcomes are resolved from the exchange's own daily high/low: target hit, stop hit, or a "
    "20-trading-day time-stop — whichever comes first (if both target and stop are touched the "
    "same day, the STOP is assumed first, so the numbers stay pessimistic, never flattering).",
    "The dashboard shows a P&L panel at the top: how many picks closed, the win rate, the "
    "average result per trade, the profit factor, and — most useful — the realized win rate "
    "PER LENS, so you can watch which lenses are actually earning their keep on live picks.",
])
why(
    "This is the measurement backbone. It turns opinions into a scoreboard: any future idea can "
    "now be judged on real, forward picks, not just on the historical backtest. It also keeps "
    "everyone honest — the ledger records the losers too, and most picks still won't reach "
    "target. Give it a few weeks of picks before reading much into the numbers."
)

# ===========================================================================
# Glossary
# ===========================================================================
doc.add_page_break()
doc.add_heading("Glossary", level=1)
glossary = [
    ("Swing trade", "Holding a stock for days to a few weeks to catch one up-move — longer "
     "than day-trading, shorter than investing."),
    ("Setup Score", "This tool's 0–100 rating of how promising a stock looks, combining "
     "relative strength, chart patterns, agreement, volume, and big-player buying."),
    ("Data-driven screener weights", "The chart-pattern part of the score. Each screener is "
     "weighted by its own MEASURED hit-rate (not by hand); consistently-losing screeners are "
     "set to zero. Lives in config/screener_weights.csv. See section 14."),
    ("EP follow-through", "A refinement to the Episodic Pivot lens: the gap must CLOSE in the "
     "top half of its day's range (it 'held'), which backtested ~19% vs ~12% for gaps that "
     "faded to the lows. A faded gap no longer earns the EP badge."),
    ("Paper ledger", "A running, self-updating book of the tool's own daily picks, with each "
     "outcome resolved against the exchange's prices — so you can see live, per-lens, how the "
     "method is actually performing. The dashboard's P&L panel reads it. See section 15."),
    ("Relative strength (vs. market)", "How much a stock has beaten the overall market "
     "(Nifty) over a recent window — here, the last 20 trading days. The strongest signal in "
     "this project."),
    ("Leadership strip / persistence", "The 20-bar row on the dashboard showing which of the "
     "last 20 days a stock was a market leader (RS ≥ +15). A full strip = a durable leader "
     "(historically the higher-win cohort); a sparse one = fresh or sporadic."),
    ("Persistent leader", "A stock that led the market on 12+ of the last 20 days — a durable "
     "trend. Backtested to hit target ~24% vs ~15% for freshly-emerged leaders."),
    ("Wyckoff SOS (Sign of Strength)", "A breakout out of a trading range on volume — "
     "'jumping the creek'. The actionable breakout entry in Wyckoff's framework."),
    ("Wyckoff LPS (Last Point of Support)", "A quiet, low-volume pullback to support AFTER a "
     "breakout — the safer, lower-chase Wyckoff entry."),
    ("Wyckoff Base (cause)", "A stock still forming its trading range (volatility "
     "contracting), no breakout yet. A watch-list, not an entry."),
    ("Episodic Pivot (EP)", "A gap-up on heavy volume — a catalyst move (earnings/news). "
     "Measures 'did something happen today', a different axis from relative strength. "
     "Backtested to ~2.5x the base rate."),
    ("Relative volume", "Today's volume divided by the stock's own recent average. 3x = "
     "trading three times its normal volume. Shown for every stock on the dashboard; it powers "
     "the Episodic Pivot and Emerging lenses."),
    ("Emerging (early move)", "A volume thrust (>=3x normal) while relative strength is rising "
     "but still modest (0 to +25) — an attempt to catch a stock waking up before it's an "
     "obvious leader. Backtested to ~1.5x the base rate: real, but about half the edge of a "
     "confirmed leader, so an earlier and lower-confidence entry."),
    ("RIGHTWAY confluence", "Cross-referencing the dashboard picks against your RIGHTWAY "
     "Telegram group's recent chatter. A confirming second opinion, not a signal — the group "
     "is promotional and often posts winners after they've moved."),
    ("Admin call (RIGHTWAY)", "A stock the RIGHTWAY broadcast channel (the admin) actually "
     "called, marked “^” on the dashboard, with a “*” when flagged important recently. Carries "
     "more weight than ordinary member mentions, but most calls in the data are 9–24 months "
     "old — treat their targets as historical context."),
    ("Backtest", "Replaying historical data to check whether a strategy would have worked, "
     "before risking real money."),
    ("Lift", "How many times better than random the tool's top picks were. 2x lift = its "
     "picks won twice as often as picking blindly."),
    ("Out-of-sample (new-date) test", "Checking the strategy on later dates it was never "
     "tuned on — the honest test that an edge is real, not a coincidence."),
    ("Hit rate / win rate", "The share of trades that reached the target."),
    ("Expected profit per trade", "The average result across many trades, counting winners "
     "and losers together. The proper way to judge a target/stop combo."),
    ("Profit factor", "Total winnings divided by total losses. Above 1.0 means profitable "
     "overall."),
    ("Take-profit (target)", "The price at which you'd sell a winning trade to lock in gains."),
    ("Stop-loss", "A pre-set price at which you sell a losing trade to cap the loss. The most "
     "important discipline in this whole system."),
    ("Reward : Risk", "How much you aim to make versus how much you risk. 5:1 means targeting "
     "a 25% gain while risking a 5% loss."),
    ("ATR (Average True Range)", "A measure of how much a stock typically moves in a day, in "
     "percent. Used to size each stock's stop-loss to its own volatility."),
    ("Volatility", "How much and how fast a stock's price swings. Jumpy stocks need wider "
     "stops so normal wiggles don't trigger them."),
    ("Chartink screener", "A saved automatic filter that returns stocks matching a chart "
     "condition (e.g. '52-week breakout')."),
    ("Delivery %", "The share of a day's volume actually taken delivery of, rather than "
     "bought and sold same-day. Higher suggests genuine accumulation."),
    ("Bulk / block deal", "A large trade reported by the exchange — a footprint of "
     "institutional (big-money) activity."),
    ("Bhavcopy", "The exchange's official end-of-day data file (prices, volumes, delivery)."),
    ("ETF / index fund", "A fund that tracks a basket or index (like NIFTYBEES). Not an "
     "individual company; excluded from this tool."),
    ("Liquidity", "How easily you can buy/sell without moving the price. We skip illiquid "
     "stocks you couldn't realistically trade."),
    ("Data gap", "A stretch where daily data wasn't collected (here, ~4 months). The tool is "
     "careful never to treat prices across a gap as consecutive."),
]
for term, d in glossary:
    p = doc.add_paragraph()
    r = p.add_run(term)
    r.bold = True
    r.font.color.rgb = ACCENT
    doc.add_paragraph(d)

# ===========================================================================
# FAQ
# ===========================================================================
doc.add_page_break()
doc.add_heading("FAQ", level=1)
faqs = [
    ("If I buy the tool's top pick, will it go up 20%?",
     "No. A high score means the stock fits a profile that historically beat the market about "
     "2–4.5x more often than random — a real edge, but most individual picks still won't reach "
     "target. It shifts the odds; it doesn't promise outcomes."),
    ("What's the best profit target, then?",
     "The target itself barely matters — a wide one (25–30%) is slightly better than 20%, but "
     "the real lever is a TIGHT stop-loss (around 5–6%) and not holding a stalled trade too "
     "long. The tool now sets a target and stop for each stock along these lines."),
    ("Why is the stop-loss different for each stock?",
     "Because a 5% drop is normal noise for a jumpy stock but a real warning for a calm one. "
     "The stop is sized to each stock's own daily volatility, so you're not shaken out of "
     "good trades or holding bad ones too long."),
    ("Why does 'buy what's already gone up' work? Isn't that backwards?",
     "It feels backwards, which is why it works — most people wait for a pullback and miss "
     "the move. Stocks that lead the market tend to keep leading for a while. That said, it "
     "does mean chasing strength, so the stop-loss matters even more."),
    ("The top pick has a big score but I'm told to reject it — how?",
     "A high score gets a name onto the list, but the LENSES decide. The two most common "
     "'high score but pass' cases: a 0/20 leadership strip (never a durable leader — a "
     "one-day flare, not a trend), and an 'ext' flag (already up >15% today — you'd be "
     "chasing). See EMAMIPAP in the worked examples. Always glance at the strip and the "
     "chase flag before the price."),
    ("What's the difference between a persistent leader and an Episodic Pivot?",
     "They're opposite profiles. A persistent leader (full leadership strip) has been strong "
     "for weeks — a proven trend. An Episodic Pivot (EP) just gapped up on volume today — a "
     "fresh catalyst with no track record required. Persistent = higher confidence, EP = "
     "earlier but less confirmed. HUHTAMAKI vs SUPREMEINF in the examples show each."),
    ("Which lens should I use day to day?",
     "Start with the top ~10 by score and the leadership strip — that's the core. Use the "
     "Wyckoff SOS preset when you want fresh breakouts, LPS for quieter pullback entries, and "
     "EP on days with earnings/news catalysts. 'Persistent leaders' is your durable-trend "
     "shortlist; 'New leaders today' surfaces fresh ideas. They're lenses on the same list, "
     "not different lists."),
    ("Why did some picks briefly show a flat 30% target instead of a stock-specific one?",
     "For a while after the data-collection gap, the volatility measure behind the per-stock "
     "stop couldn't be computed, so the tool fell back to a flat 6% stop / 30% target. That's "
     "resolved: the gap was backfilled from the exchange's archives and history now runs "
     "unbroken from July 2024, so every stock gets its own volatility-sized levels "
     "immediately. If you ever see the '[flat, no ATR]' tag again, it means a stock is "
     "genuinely new to the exchange or data collection has lapsed — rerun "
     "backfill_bhavcopy.py in that case."),
    ("Some picks are 'BE' or Trade-to-Trade stocks — are those safe to trade?",
     "They're tradeable (delivery-only settlement, which is fine for a swing trade you hold "
     "anyway), and we include them so nothing in your watchlist gets silently dropped. But "
     "they tend to be smaller and more volatile, and the 'BZ' ones are under exchange "
     "surveillance with restricted price bands. Their backtested numbers look better than "
     "you'd realistically capture, so give them extra chart scrutiny and don't over-size them."),
    ("Should I buy a stock just because my RIGHTWAY Telegram group is calling it?",
     "No. The group is heavily promotional and often posts winners after they've already "
     "moved, and about 43 of 60 dashboard names get mentioned on any given day — so a mention "
     "isn't selective. Use it only as a confirming nudge on a stock the tool ALREADY rates "
     "(especially 'strong setup + fresh ADMIN call'), never as a standalone reason. The admin "
     "'^'/'*' markers matter far more than raw mention count, and most admin targets in the "
     "data are 9-24 months old — historical context, not live levels. The tool's own "
     "volatility-based target and stop are the only ones with a backtest behind them."),
    ("What is 'Emerging' and should I prefer it over a confirmed leader?",
     "Emerging catches a stock a step earlier — a volume thrust (3x+ normal) while its "
     "relative strength is rising but still modest. It's tempting because you're early, but "
     "testing showed it hits target at ~1.5x the base rate versus ~3x for a confirmed leader: "
     "an earlier, LOWER-confidence entry, not a better one. Use it for a smaller, speculative "
     "position, and let the confirmed-leader cohort remain your core."),
    ("How often should I run this?",
     "Once per trading day after the market closes and the daily sheet is built. The signals "
     "are about what changed that day, so re-running on stale data adds nothing. If you use "
     "the RIGHTWAY confluence, run telegram_fetch.py and parse_telegram_recs.py at end of day "
     "too, before make_dashboard.py."),
    ("Should I trust this over a fancy AI/machine-learning model?",
     "For now, yes. An earlier machine-learning attempt on this exact data failed — winners "
     "are too rare, so it just learned to predict 'no'. The simple, transparent scoring here "
     "beat it, and the relative-strength discovery gave a bigger improvement than the model "
     "ever did."),
    ("What could still go wrong / what's not accounted for?",
     "Trading costs and slippage aren't included, results use daily closing prices only, and "
     "everything is measured over a single ~13-month window. Treat the findings as a solid "
     "direction, not precise guarantees — and never skip the stop-loss."),
]
for q, a in faqs:
    faq(q, a)

doc.add_paragraph()
foot = doc.add_paragraph()
fr = foot.add_run(
    "Built and tested on the Odin's Watchlist project. Historical tests span 2,554 stocks "
    "over 158 trading days (Jul 2025 – Jul 2026) for the watchlist, and 2,992 stocks (EQ + BE "
    "+ BZ series) over 528 unbroken trading days (Jul 2024 – Jul 2026) for price history. "
    "Findings are directional, not guarantees; use a stop-loss."
)
fr.italic = True
fr.font.size = Pt(9)
fr.font.color.rgb = INK_SOFT

OUT = "/Users/vinaykusuma/Documents/AlgoTrade/odins_watchlist/working_version/Odins_Watchlist_Guide.docx"
doc.save(OUT)
print("Saved:", OUT)
