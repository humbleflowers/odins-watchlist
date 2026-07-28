"""Generate the Odin's Watchlist WORKED-EXAMPLES document.

For each use case: the real, dated trades that would have resulted, pulled
straight from the backtest panels (examples.json produced by extract_examples.py).
Every row is a genuine (symbol, date) with its real recorded outcome:
  WIN  = the +25% target was hit BEFORE the -5% stop within 30 sessions.
  LOSS = the stop hit / the trade expired before target.
  peak = the highest the stock actually reached within 30 sessions (%).
Nothing is fabricated. Rates shown are measured on this 12-month window.
"""
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SC = Path(__file__).resolve().parent
DATA = json.loads((SC / "examples.json").read_text())
META = DATA["_meta"]

ACCENT = RGBColor(0x7A, 0x54, 0x17)
INK_SOFT = RGBColor(0x55, 0x52, 0x4A)
POSITIVE = RGBColor(0x2E, 0x6B, 0x4C)
WARN = RGBColor(0x9C, 0x3E, 0x2E)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
for i, size in [(1, 20), (2, 15), (3, 12.5)]:
    h = doc.styles[f"Heading {i}"]
    h.font.size = Pt(size)
    h.font.color.rgb = RGBColor(0x1C, 0x1B, 0x17)
    h.font.bold = True


def shade(cell, hexc):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hexc)
    cell._tc.get_or_add_tcPr().append(shd)


def set_width(cell, inches):
    cell.width = Inches(inches)


def callout(title, body, kind="info"):
    color = {"info": ACCENT, "good": POSITIVE, "warn": WARN}[kind]
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22)
    r = p.add_run(title.upper()); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = color
    b = doc.add_paragraph(body)
    b.paragraph_format.left_indent = Inches(0.22)
    b.paragraph_format.space_after = Pt(12)


PENNY = {"EXCEL"}  # drop sub-Rs2 illiquid names we wouldn't trade anyway


def examples_table(key, metric_label=None, metric_key=None):
    d = DATA[key]
    rows = [e for e in d["examples"] if e["symbol"] not in PENNY]
    headers = ["Date", "Stock", "Entry", "Score", "RS"]
    if metric_key:
        headers.append(metric_label)
    headers += ["Result", "Peak in 30d"]
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 5"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, htxt in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = htxt
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True; r.font.size = Pt(9.5)
        shade(c, "EFE2C6")
    for e in rows:
        cells = t.add_row().cells
        vals = [e["date"], e["symbol"], f"Rs{e['price']:g}", str(e["score"]), f"{e['rs']:+d}"]
        if metric_key:
            mv = e["extra"].get(metric_key, "")
            vals.append(str(mv))
        result = "WIN - target hit" if e["win"] else "LOSS - stopped"
        peak = f"+{e['peak']}%" if (e["peak"] is not None and e["peak"] >= 0) else (f"{e['peak']}%" if e["peak"] is not None else "n/a")
        vals += [result, peak]
        for i, v in enumerate(vals):
            cell = cells[i]; cell.text = v
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
                    if headers[i] == "Result":
                        r.font.color.rgb = POSITIVE if e["win"] else WARN
                        r.bold = True
    doc.add_paragraph()


def rate_line(key):
    d = DATA[key]
    p = doc.add_paragraph()
    r = p.add_run("Measured on this data: ")
    r.bold = True; r.font.color.rgb = ACCENT
    p.add_run(
        f"{d['n']} completed trades in this cohort. {d['win_rate']}% hit the +25% target before "
        f"the -5% stop; {d['touch_rate']}% touched +20% at some point. "
        f"(Whole-watchlist base rate: {META['base_win']}% / {META['base_touch']}%.)"
    )


# ===========================================================================
doc.add_heading("Odin's Watchlist - Worked Examples by Use Case", level=0)
sub = doc.add_paragraph()
sr = sub.add_run("Real, dated trades each use case would have produced - with the outcome that "
                 "actually happened")
sr.italic = True; sr.font.size = Pt(13)
doc.add_paragraph()
doc.add_paragraph(
    "This companion to the Use-Cases playbook answers one question: 'if I had traded each of "
    f"these patterns in the past, what would have happened?' Every trade below is a genuine "
    f"stock on a genuine date from the backtest window ({META['dates']}, {META['ndates']} "
    f"trading days, {META['n_resolved']:,} completed trades). The result column is the outcome "
    "the data actually recorded, using the tool's own exit rule: a +25% target against a -5% "
    "stop, first-hit within 30 trading sessions."
)
callout(
    "How to read the results honestly",
    "WIN means the +25% target was reached BEFORE the -5% stop. LOSS means the stop hit or the "
    "trade expired first. 'Peak in 30d' is the highest the stock actually got to - a LOSS can "
    "still show a small positive peak if it rose a little, then fell to the stop before reaching "
    "target. Each block shows a few winners AND a few losers on purpose: even the best use "
    "cases win only about a quarter of the time, so the losers are the normal texture of the "
    "strategy, not a malfunction. Numbers are one 12-month window on daily closes before costs "
    "- directional, not a promise.",
    kind="warn",
)
callout(
    "The single most important pattern in this whole document",
    f"The whole-watchlist base rate is just {META['base_win']}% (target hit before stop). The "
    "core top-10 by score lifts that to 20%, and the two strongest lenses - persistent "
    "leadership and Episodic Pivot - push it to ~24-27% touched. That is the entire edge: "
    "roughly 4-5x better than random, NOT a high win rate. You make money by taking many of "
    "these with a tight stop, letting the ~1-in-4 winners run to a wide target.",
    kind="good",
)
callout(
    "Two upgrades added since these examples were drawn",
    "A later round of feature backtests produced two changes that make the picks a little "
    "cleaner than the raw examples below: (1) the chart-pattern score is now DATA-DRIVEN - each "
    "screener weighted by its measured hit-rate, dead screeners zeroed; and (2) the Episodic "
    "Pivot badge now fires only when the gap HELD (closed top-half of range), which backtested "
    "~19% vs ~12% for faded gaps. Several other ideas (regime filter, sector rotation, "
    "delivery-trend, smart-money deals, trailing stops) were tested and REJECTED - see section "
    "14 of the main Guide. A paper ledger now also tracks these picks live, per lens.",
    kind="info",
)
doc.add_page_break()

# ===========================================================================
doc.add_heading("Tier 1 - the patterns with a clear edge", level=1)
doc.add_paragraph(
    "These three beat the whole-market base rate by 4-5x AND stand up as the strongest cohorts "
    "in testing. This is where the money is."
)

doc.add_heading("Use case 1 - Core top-10 by Setup Score", level=2)
doc.add_paragraph("The daily staple: buy from the ten highest scores. Two big winners, two "
                  "high-score losers - note the losers were BOTH score-100 names that still "
                  "failed, which is exactly why you need the stop and can't skip the lenses.")
rate_line("uc1"); examples_table("uc1")

doc.add_heading("Use case 2 - Persistent leaders (durable trend)", level=2)
doc.add_paragraph("Stocks that led the market on 12+ of the last 20 days. STALLION here is the "
                  "textbook case - a full 20/20 strip that ran +88%. The losers (UNICHEMLAB, "
                  "WABAG) show even proven leaders fail one trade in four-ish; the stop caps it.")
rate_line("uc2"); examples_table("uc2", "20d strip", "strip")
callout(
    "Persistence is the standout - the numbers that matter",
    "Measured WITHIN each day's top-10 by score, persistent leaders (strip >=12) hit the target "
    "26% of the time versus 15% for freshly-emerged leaders - the single biggest reliable "
    "upgrade any lens gives you. Note also that a persistent leader ALSO firing a fresh breakout "
    "the same day is intuitively the premium setup, but in this particular window that exact "
    "intersection was too small a sample (49 trades, mostly clustered in one choppy stretch) to "
    "claim a clean number - so treat 'durable leader + breakout' as a quality preference, and "
    "lean on the persistence number itself, which is robust.",
    kind="good",
)

doc.add_heading("Use case 4 - Episodic Pivot (catalyst gap-up)", level=2)
doc.add_paragraph("A gap of 4%+ on 3x+ volume - a catalyst. SABTNL gapped and ran +80%. Notice "
                  "these fire even from modest scores (43-50): the gap is a DIFFERENT signal "
                  "from strength, which is why it adds something. The losers gapped but stalled.")
rate_line("uc4"); examples_table("uc4", "gap% / rvol", "rvol")
callout(
    "Why EP earns its place",
    "At 27% touched it is the highest touch-rate cohort here, and it is independent of the "
    "score - it answers 'did something happen today', not 'is it already strong'. Because "
    "catalyst moves can fade fast, the tight stop matters even more than usual.",
    kind="good",
)
doc.add_page_break()

# ===========================================================================
doc.add_heading("Tier 2 - situational lenses (use to CHOOSE, not to replace the score)", level=1)
doc.add_paragraph(
    "These next patterns, applied across the shortlist, won at ~11-14% - BELOW the score-top-10's "
    "20%. That is not because they're bad, but because the cohorts include lower-scored names. "
    "The lesson matches the whole project: relative strength (the score) is the edge; these "
    "lenses help you pick a good ENTRY among already-strong names - they are not a substitute "
    "for a high score. The examples show real winners, but also why you shouldn't lead with them."
)

doc.add_heading("Use case 5 - Wyckoff SOS / breakout trigger", level=2)
doc.add_paragraph("Breakout-structure entries. Big winners exist (MANAKALUCO +69%, SILVERBEES "
                  "+68%) - but the two score-100 losers are the same AARTECH / DIGITIDE that "
                  "broke out and immediately failed. A breakout without durable leadership is "
                  "a coin-flip; confirm the strip first.")
rate_line("uc5"); examples_table("uc5", "phase", "phase")

doc.add_heading("Use case 6 - Wyckoff LPS / rebound pullback", level=2)
doc.add_paragraph("Buying the quiet pullback to support after strength. Winners like TCIFINANCE "
                  "(+63%) came off very high RS. Best used to get a cheaper, tighter-stop entry "
                  "into a name you already rate - not as a stand-alone screen.")
rate_line("uc6"); examples_table("uc6", "phase", "phase")

doc.add_heading("Use case 7 - Emerging (early move, before leadership)", level=2)
doc.add_paragraph("A volume thrust (3x+) while RS is still modest (0 to +25) - catching the "
                  "wake-up. Real early winners exist (SHARDUL +68%, DPABHUSHAN +37%), but at just "
                  "4% target-hit / 7% touched this is barely above the base rate. This is the "
                  "honest confirmation of the guide's warning: earlier, but LOWER confidence - a "
                  "small speculative satellite, never your core.")
rate_line("uc7"); examples_table("uc7", "rvol", "rvol")
doc.add_page_break()

# ===========================================================================
doc.add_heading("Anti-patterns - what these examples warn you AWAY from", level=1)

doc.add_heading("Trap 1 - high score, empty leadership strip (the one-day flare)", level=2)
doc.add_paragraph("Stocks scoring 75+ but with a near-empty 20-day strip (never a durable "
                  "leader). At just 10% target-hit - HALF the clean top-10's 20% - this cohort "
                  "is where high scores go to disappoint. DIGITIDE / GODREJIND / MASTERTR all "
                  "scored 90+ and still lost. The strip is the single best filter that saves you.")
rate_line("anti1"); examples_table("anti1", "20d strip", "strip")

doc.add_heading("Trap 2 - high delivery % (the lens we tested and dropped)", level=2)
doc.add_paragraph("High-delivery names that ALSO made the shortlist won ~26% - which looks "
                  "great, until you realise that's the SAME rate as any other strong shortlist "
                  "name. That's the whole point: delivery % added nothing BEYOND the strength "
                  "that already got them listed, so as a stand-alone screen it's noise. It was "
                  "measured, found redundant, and left out of the score.")
rate_line("anti2"); examples_table("anti2", "deliv%", "deliv%")

callout(
    "The thread through every example",
    "The SCORE (relative strength) is the edge - it lifts you from a 5% base to a 20% hit rate. "
    "PERSISTENCE and EPISODIC PIVOT are the two lenses that genuinely add more. Everything else "
    "helps you choose an entry or avoid a trap, but cannot replace a high score plus a durable "
    "leadership strip. And in every single cohort - even the best - most trades still lose, so "
    "the tight stop and the wide target are what convert this edge into money.",
    kind="good",
)

doc.add_paragraph()
foot = doc.add_paragraph()
fr = foot.add_run(
    f"All trades extracted from the project's backtest panels (labeled_panel.csv / "
    f"ohlc_indicator_panel.csv), window {META['dates']}, {META['ndates']} trading days, "
    f"{META['n_resolved']:,} completed trades. Exit rule: +25% target / -5% stop, first-hit "
    "within 30 sessions. Daily closes, before costs; sub-Rs2 illiquid names excluded. "
    "Real outcomes, one market window - directional, not guarantees. Always use a stop-loss."
)
fr.italic = True; fr.font.size = Pt(9); fr.font.color.rgb = INK_SOFT

OUT = "/Users/vinaykusuma/Documents/AlgoTrade/odins_watchlist/working_version/Odins_Watchlist_UseCase_Examples.docx"
doc.save(OUT)
print("Saved:", OUT)
