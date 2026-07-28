"""Generate a plain-English BEGINNER'S guide to swing trading for Odin's
Watchlist -- written for someone brand new to the stock market. No jargon
assumed; every term is explained the first time it appears, with a glossary
and FAQ at the end. Grounded in the same backtested findings as the main Guide.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x7A, 0x54, 0x17)
INK_SOFT = RGBColor(0x55, 0x52, 0x4A)
POSITIVE = RGBColor(0x2E, 0x6B, 0x4C)
WARN = RGBColor(0x9C, 0x3E, 0x2E)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
for i, size in [(1, 20), (2, 15), (3, 12.5)]:
    h = doc.styles[f"Heading {i}"]
    h.font.size = Pt(size)
    h.font.color.rgb = RGBColor(0x1C, 0x1B, 0x17)
    h.font.bold = True


def shade(cell, hexc):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hexc)
    cell._tc.get_or_add_tcPr().append(shd)


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 5"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, htxt in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = htxt
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
        shade(c, "EFE2C6")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    doc.add_paragraph()


def bullets(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def numbered(items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def callout(title, body, kind="info"):
    color = {"info": ACCENT, "good": POSITIVE, "warn": WARN}[kind]
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22)
    r = p.add_run(title.upper())
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = color
    b = doc.add_paragraph(body)
    b.paragraph_format.left_indent = Inches(0.22)
    b.paragraph_format.space_after = Pt(12)


def plain(term, meaning):
    """A 'in plain words' inline definition."""
    p = doc.add_paragraph()
    r = p.add_run(term + " — ")
    r.bold = True
    r.font.color.rgb = ACCENT
    p.add_run(meaning)


def faq(q, a):
    p = doc.add_paragraph()
    p.add_run(q).bold = True
    doc.add_paragraph(a)


# ===========================================================================
doc.add_heading("Odin's Watchlist — A Beginner's Guide", level=0)
sub = doc.add_paragraph()
sr = sub.add_run("Brand new to the stock market? Start here. Every idea explained in plain "
                 "words, then the handful of situations actually worth trading.")
sr.italic = True
sr.font.size = Pt(13)
doc.add_paragraph()
intro = doc.add_paragraph()
intro.add_run(
    "This guide assumes you know nothing about the stock market and explains everything from "
    "the ground up. Read it once, top to bottom. By the end you'll understand what a stock is, "
    "how people make (and lose) money trading, the small number of ideas that actually matter, "
    "and — most importantly — the two or three situations that are genuinely worth trading and "
    "the traps to avoid. There's a plain-English dictionary at the very end."
)
callout(
    "The one promise, and the one warning",
    "PROMISE: you do not need maths, finance jargon, or a fancy screen to follow this. "
    "WARNING: trading real money is risky and most beginners lose money by rushing. Read the "
    "whole guide, practise WITHOUT real money first (that's called 'paper trading'), and never "
    "risk money you can't afford to lose. Nothing here is financial advice — it's an "
    "explanation of how the tool works.",
    kind="warn",
)
doc.add_page_break()

# ===========================================================================
doc.add_heading("1. What is the stock market, really?", level=1)
doc.add_paragraph(
    "A COMPANY that wants money to grow can sell small pieces of itself to the public. Each "
    "small piece is a SHARE (also called a STOCK). If you own a share, you own a tiny slice of "
    "that company. There are thousands of companies you can buy shares in."
)
plain("Share price", "what one share costs right now. It goes up and down every second the "
      "market is open, based purely on supply and demand: if more people want to BUY a share "
      "than SELL it, the price rises; if more want to sell, it falls. That's the whole engine.")
plain("Stock exchange", "the marketplace where all this buying and selling happens. In India "
      "the main one is the NSE (National Stock Exchange). It's open on weekdays, roughly 9:15am "
      "to 3:30pm.")
plain("Broker / trading app", "the service you use to place buy and sell orders (Zerodha, "
      "Groww, etc.). You put money in, and it lets you buy and sell shares.")
plain("Nifty", "a single number that tracks the overall Indian market (an average of 50 big "
      "companies). When people say 'the market went up today', they usually mean Nifty went up. "
      "It's the yardstick everything else is measured against.")
doc.add_paragraph(
    "That's it. You buy shares hoping to sell them later at a higher price. The difference "
    "(minus small fees) is your profit — or, if the price fell, your loss."
)

# ===========================================================================
doc.add_heading("2. Three ways people trade — and where this tool fits", level=1)
table(
    ["Style", "How long you hold", "In plain words"],
    [
        ["Investing", "Years", "Buy good companies and hold for a long time. Slow and steady."],
        ["Swing trading", "A few days to a few weeks", "Try to catch ONE up-move in a stock, "
         "then sell. This is what Odin's Watchlist is built for."],
        ["Day trading", "Minutes to hours", "Buy and sell the same day. Fast, stressful, and "
         "the riskiest — NOT what this tool does."],
    ],
)
callout(
    "What this tool actually is",
    "Odin's Watchlist is a SWING-TRADING assistant. Every day it scans thousands of Indian "
    "stocks and hands you a short list of the ones that, historically, had the best chance of "
    "jumping over the following days-to-weeks. It does NOT tell you the future — it tips the "
    "odds slightly in your favour and shows you a sensible plan for each stock. You still "
    "decide, and you still take the risk.",
    kind="info",
)

# ===========================================================================
doc.add_heading("3. The five ideas you must understand (that's all)", level=1)
doc.add_paragraph("Ignore the hundreds of terms out there. These five cover almost everything:")

doc.add_heading("① Percent change (% change)", level=3)
doc.add_paragraph(
    "How much a price moved, as a percentage. If a Rs 100 stock rises to Rs 110, that's +10%. "
    "Percentages let you compare a Rs 50 stock and a Rs 5,000 stock fairly."
)
doc.add_heading("② Volume", level=3)
doc.add_paragraph(
    "How MANY shares were bought and sold today. High volume means lots of people are "
    "interested — the move is 'real'. A price rise on tiny volume is weak and often fades. "
    "Think of volume as the crowd size behind a move."
)
doc.add_heading("③ Trend", level=3)
doc.add_paragraph(
    "The general direction over recent weeks — up, down, or sideways. The single most useful "
    "habit in trading is to favour stocks in an UP trend and avoid ones in a DOWN trend. "
    "'The trend is your friend' is a cliché because it's true."
)
doc.add_heading("④ Support and resistance (the floor and the ceiling)", level=3)
doc.add_paragraph(
    "As a stock trades, certain price levels keep stopping it. A SUPPORT level acts like a "
    "FLOOR — the price keeps bouncing UP off it. A RESISTANCE level acts like a CEILING — the "
    "price keeps getting pushed DOWN from it. These aren't magic; they're just prices where "
    "lots of buyers or sellers have shown up before."
)
doc.add_heading("⑤ Breakout", level=3)
doc.add_paragraph(
    "When a price finally pushes THROUGH a ceiling (resistance) it had been stuck under, that's "
    "a BREAKOUT. It often means the buyers have taken control and the stock may run higher. "
    "Breakouts — especially on high volume — are one of the main things this tool looks for."
)
callout(
    "A picture in words",
    "Imagine a stock bouncing inside a box for weeks: it keeps falling from Rs 100 (the "
    "ceiling) and bouncing off Rs 90 (the floor). One day heavy buying pushes it to Rs 104 on "
    "big volume — it 'broke out' of the box. That break, with a crowd (volume) behind it, is "
    "the kind of moment a swing trader wants to catch.",
    kind="good",
)

# ===========================================================================
doc.add_heading("4. The two rules that protect your money", level=1)
doc.add_paragraph(
    "Before any 'best stock to buy' talk, learn these two. They matter more than picking "
    "winners. Beginners who skip them blow up their accounts."
)
doc.add_heading("Rule 1 — Always set a stop-loss", level=3)
plain("Stop-loss", "a price you decide IN ADVANCE at which you'll sell if the trade goes "
      "against you, to cap your loss. It's an ejector seat. If you buy at Rs 100 and set a "
      "stop-loss at Rs 95, you've decided the most you'll lose is about 5%. Set it BEFORE you "
      "buy, and actually honour it — don't move it lower hoping the stock recovers.")
doc.add_paragraph(
    "Why it's non-negotiable: a few small, controlled losses are survivable. One big loss where "
    "you 'held on hoping' can wipe out many good trades. The tool prints a suggested stop-loss "
    "for every stock — use it."
)
doc.add_heading("Rule 2 — Only risk a tiny slice per trade (position sizing)", level=3)
plain("Position sizing", "deciding HOW MANY shares to buy so that if your stop-loss is hit, "
      "you only lose a small, fixed slice of your money — usually 1% or 2%.")
doc.add_paragraph("Worked example, in plain numbers:")
bullets([
    "Say your total trading money is Rs 2,00,000 and you decide to risk 1% per trade = Rs 2,000 "
    "maximum loss on any single trade.",
    "You want to buy a stock at Rs 100 with a stop-loss at Rs 95 — so you'd lose Rs 5 per share "
    "if stopped out.",
    "Rs 2,000 max loss ÷ Rs 5 per share = 400 shares. So you buy 400 shares (Rs 40,000 worth). "
    "If it hits the stop, you lose Rs 2,000 — exactly your 1%. No surprises.",
])
callout(
    "The dashboard does this maths for you",
    "Type your total capital and your chosen risk-% once into the dashboard's top bar, and "
    "every stock row instantly shows how many shares to buy. You never have to calculate it by "
    "hand. Sizing every trade to the same small risk is what lets you survive a losing streak — "
    "and losing streaks are normal, even for good strategies.",
    kind="good",
)

# ===========================================================================
doc.add_heading("5. The best cases to trade — what actually works", level=1)
doc.add_paragraph(
    "This is the heart of it. The tool was tested on more than a year of real history to find "
    "what genuinely improves your odds. A few patterns stood out. Here they are in plain words, "
    "strongest first. You don't trade EVERYTHING on the list — you wait for these."
)

doc.add_heading("Best case 1 — A proven leader that breaks out today", level=2)
doc.add_paragraph(
    "The single strongest setup. Two things line up at once: (a) the stock has been one of the "
    "market's strongest performers for weeks (a 'proven leader'), AND (b) it breaks out of its "
    "box TODAY on good volume."
)
plain("Relative strength (why 'leader' matters)", "how much a stock has beaten the overall "
      "market (Nifty) recently. A stock that keeps out-running the market is showing that big, "
      "informed buyers want it. Surprisingly, buying stocks that are ALREADY strong works better "
      "than hunting for 'cheap' beaten-down stocks — because strength tends to continue.")
bullets([
    "What to look for on the dashboard: a high Score, a nearly-full 20-day 'leadership strip' "
    "(the little row of bars showing it's been a leader most days), and a fresh breakout badge "
    "today.",
    "Why it's the best: in testing, these 'durable leader' names hit their target about 1 in 4 "
    "times — versus about 1 in 7 for the average pick. That's a meaningful edge.",
    "The catch: if it's ALREADY jumped a lot today (say up 15%+), you're late — wait for it to "
    "calm down rather than chasing.",
])

doc.add_heading("Best case 2 — A catalyst gap that holds", level=2)
doc.add_paragraph(
    "Sometimes news (good results, a big order) makes a stock GAP UP — it opens much higher than "
    "yesterday's close — on huge volume. If that gap HOLDS (the stock stays strong into the "
    "close rather than fading back), it often keeps running. The tool calls this an "
    "'Episodic Pivot'."
)
bullets([
    "What to look for: the EP badge on the dashboard (a gap of 4%+ on 3x+ normal volume that "
    "closed in the top half of the day's range).",
    "Why it works: the gap means 'something real just happened' — brand-new information the "
    "usual signals don't capture. Tested, these hit target more often than the average pick.",
    "The catch: gaps can fade fast. Keep the position smaller and honour the stop tightly.",
])

doc.add_heading("Best case 3 — Simply: buy strength, near a fresh trigger", level=2)
doc.add_paragraph(
    "If cases 1 and 2 feel like a lot, the simple version is: from the day's top ten by Score, "
    "prefer names with strong relative strength that are breaking out today on volume, and skip "
    "the rest. That top slice is where the edge lives."
)
table(
    ["Green lights (want to see)", "Red lights (walk away)"],
    [
        ["Strong relative strength (a leader)", "Lagging the market (weak)"],
        ["A fresh breakout today, on volume", "Already up a lot today (you're chasing)"],
        ["A full leadership strip (durable)", "An empty strip (a one-day flare, no history)"],
        ["A clean chart you understand", "A messy chart, or you 'heard a tip'"],
    ],
)

# ===========================================================================
doc.add_heading("6. The traps — what to AVOID", level=1)
doc.add_paragraph("Most beginner losses come from a short list of avoidable mistakes:")
bullets([
    "CHASING. Buying a stock that's already shot up 15-20% today because you're afraid to miss "
    "out. You're usually buying right before it pulls back. Wait, or skip it.",
    "THE ONE-DAY FLARE. A stock with a big Score today but NO history of leadership (an empty "
    "20-day strip). It spiked once; there's no trend underneath. High score is necessary, not "
    "sufficient — always glance at the strip.",
    "TINY PROFIT TARGETS. Selling for a quick +5% feels safe but loses money over time, because "
    "you're risking a 5% stop to make 5% — the maths doesn't work. Aim for wider moves and let "
    "winners run.",
    "SKIPPING THE STOP-LOSS. The fastest way to turn a small loss into a huge one. Never.",
    "TIP-CHASING / hype. Buying because a Telegram group, YouTuber, or friend said so. The tool "
    "even shows what a trading group is discussing — but only as a weak 'second opinion', never "
    "a reason to buy on its own (those channels often hype a stock AFTER it has already moved).",
])

# ===========================================================================
doc.add_heading("7. Your simple daily routine with the tool", level=1)
numbered([
    "After the market closes, open the dashboard (a single web page). At the top you'll see the "
    "day's shortlist, already scored and ranked.",
    "Type your capital and risk-% (e.g. 1%) into the top bar once, so every row shows how many "
    "shares to buy.",
    "Start with the top ~10 names. Look for Best Case 1 (proven leader + fresh breakout) and "
    "Best Case 2 (a holding gap). Star the ones you like.",
    "For each starred name, click the 'TV' link to see its chart on TradingView and confirm it "
    "looks like a clean breakout you understand — never buy a chart that confuses you.",
    "Skip anything with the red lights from section 5 (chasing, empty strip, lagging).",
    "If you take a trade: set the printed STOP-LOSS before you buy, buy the share count the "
    "sizer showed, and let the TARGET run. Don't babysit it hour by hour.",
    "Track how you're doing over time — the dashboard's paper-ledger panel shows how the tool's "
    "own picks have performed, so you can build trust before risking much.",
])
callout(
    "Practise first — 'paper trading'",
    "For your first few weeks, don't use real money. Write down the trades you WOULD take (entry, "
    "stop, target, share count) and check a week later what happened. This 'paper trading' "
    "teaches you the process and the emotions for free. Only add real money once you can follow "
    "the rules calmly.",
    kind="good",
)

# ===========================================================================
doc.add_heading("8. Honest expectations (please read twice)", level=1)
bullets([
    "MOST individual picks will NOT hit their target. Even the best setups work only about 1 "
    "time in 4. That is normal and still profitable IF you cut losers quickly and let the few "
    "winners run far.",
    "You WILL have losing streaks. Small, controlled losses in a row are the normal path to the "
    "occasional big winner. This is exactly why position sizing (Rule 2) matters.",
    "The tool improves your ODDS versus picking blindly — it does not remove risk or predict "
    "the future. Anyone promising certainty is lying.",
    "Results in testing use end-of-day prices and ignore fees, so real returns are a bit lower. "
    "Treat every number as a helpful guide, not a guarantee.",
    "This is educational software, not financial advice. Start tiny, protect your money, and "
    "never trade money you need for real life.",
])

# ===========================================================================
doc.add_page_break()
doc.add_heading("Plain-English dictionary", level=1)
glossary = [
    ("Share / Stock", "A tiny ownership piece of a company that you can buy and sell."),
    ("Share price", "What one share costs right now; moves with supply and demand."),
    ("Stock exchange (NSE)", "The marketplace where shares are bought and sold."),
    ("Broker / trading app", "The service you use to place buy/sell orders and hold your shares."),
    ("Nifty", "A single number tracking 50 big Indian companies — a stand-in for 'the market'."),
    ("Swing trade", "Holding a stock for days to a few weeks to catch one up-move."),
    ("% change", "How much a price moved, in percent (Rs 100 to Rs 110 = +10%)."),
    ("Volume", "How many shares traded — the 'crowd size' behind a price move."),
    ("Trend", "The general direction (up / down / sideways) over recent weeks."),
    ("Support (floor)", "A price level the stock keeps bouncing UP from."),
    ("Resistance (ceiling)", "A price level the stock keeps getting pushed DOWN from."),
    ("Breakout", "When the price pushes THROUGH a resistance ceiling — often a buy trigger, "
     "especially on high volume."),
    ("Relative strength", "How much a stock has beaten the overall market recently. High = a "
     "'leader'. The strongest signal this tool uses."),
    ("Leadership strip", "The little 20-bar row on the dashboard showing how many of the last "
     "20 days the stock was a market leader. Full = durable; empty = a one-day flare."),
    ("Gap up", "When a stock opens much higher than yesterday's close, usually on news."),
    ("Episodic Pivot (EP)", "A gap-up on heavy volume that holds — a fresh catalyst move."),
    ("Stop-loss", "A pre-set price where you sell to cap a loss. The most important discipline "
     "in trading."),
    ("Target (take-profit)", "The price where you'd sell a winning trade to lock in the gain."),
    ("Position sizing", "Choosing how many shares to buy so a stop-out only costs a small, "
     "fixed slice (e.g. 1%) of your money."),
    ("Reward : risk", "How much you aim to make versus how much you risk. 5:1 means targeting a "
     "25% gain while risking a 5% loss."),
    ("Volatility", "How much and how fast a stock's price swings. Jumpy stocks need wider stops."),
    ("Score", "The tool's 0-100 rating of how promising a stock looks today. Higher = better, "
     "but still not a promise."),
    ("Paper trading", "Practising trades on paper (no real money) to learn the process safely."),
    ("Chasing", "Buying a stock that's already jumped a lot today — usually right before it "
     "pulls back. A classic beginner mistake."),
]
for term, d in glossary:
    p = doc.add_paragraph()
    r = p.add_run(term)
    r.bold = True
    r.font.color.rgb = ACCENT
    doc.add_paragraph(d)

# ===========================================================================
doc.add_page_break()
doc.add_heading("Beginner FAQ", level=1)
faqs = [
    ("How much money do I need to start?",
     "Less than you think, but the real answer is: start with an amount you could lose entirely "
     "without it affecting your life, and practise on paper first. The habits matter far more "
     "than the amount."),
    ("If a stock has a high Score, will it go up?",
     "No. A high Score means it fits a profile that historically beat the market a few times "
     "more often than random — a real edge, but most single picks still don't reach target. It "
     "shifts the odds; it never promises an outcome."),
    ("Why buy a stock that's already gone up? Isn't cheap better?",
     "It feels backwards, but strength tends to continue — stocks leading the market often keep "
     "leading for a while, while 'cheap' beaten-down stocks are often cheap for a reason. That "
     "said, don't chase a stock that has already spiked hard TODAY; wait for a calmer entry."),
    ("What's the most common beginner mistake?",
     "Two, tied: skipping the stop-loss (turning small losses into huge ones), and chasing a "
     "stock that has already run up a lot today. Fix those two and you're ahead of most "
     "beginners."),
    ("How many stocks should I buy at once?",
     "Enough to spread your risk (a few names), each sized to the same small risk, but not so "
     "many you can't keep track. Quality over quantity — a couple of the best setups beats ten "
     "mediocre ones."),
    ("Should I trust tips from Telegram groups or YouTube?",
     "Treat them as noise until proven otherwise. This tool can show what a trading group is "
     "discussing, but only as a weak confirming nudge on a stock it already rates — never as a "
     "standalone reason. Promotional channels often hype stocks after they've already moved."),
    ("How long do I hold a swing trade?",
     "Days to a few weeks. You exit when it hits your target, hits your stop-loss, or has "
     "clearly stalled for a couple of weeks. You don't hold forever hoping."),
    ("Is this financial advice?",
     "No. It's an educational tool that improves your odds and enforces good habits. You make "
     "every decision and take every risk. Start small and protect your capital."),
]
for q, a in faqs:
    faq(q, a)

doc.add_paragraph()
foot = doc.add_paragraph()
fr = foot.add_run(
    "A beginner's companion to Odin's Watchlist. Once these ideas feel comfortable, read the "
    "main Guide and the Use-Cases playbook for the deeper detail. Trading is risky; practise on "
    "paper, always use a stop-loss, and never risk money you can't afford to lose. Not financial "
    "advice."
)
fr.italic = True
fr.font.size = Pt(9)
fr.font.color.rgb = INK_SOFT

OUT = "/Users/vinaykusuma/Documents/AlgoTrade/odins_watchlist/working_version/Odins_Watchlist_Beginner_Guide.docx"
doc.save(OUT)
print("Saved:", OUT)
