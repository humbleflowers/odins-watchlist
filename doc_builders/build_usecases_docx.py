"""Generate the Odin's Watchlist USE-CASES / PLAYBOOK document.

Each use case = one tradeable pattern the dashboard can isolate, presented with:
  - what it is / how to isolate it,
  - the historical backtest finding (hit rate vs base rate, lift),
  - the expected gain (per-trade expectancy, honestly labelled),
  - exactly how to trade it (entry / target / stop / hold / sizing).

All numbers come from the project's own backtests (find_swing_candidates.py,
backtest_swing_candidates.py, optimize_target.py) as summarised in
Odins_Watchlist_Guide.docx. They are directional: one ~13-month window, daily
closes, no costs/slippage. Treat the DIRECTION as robust, exact % as approximate.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x7A, 0x54, 0x17)
INK_SOFT = RGBColor(0x55, 0x52, 0x4A)
POSITIVE = RGBColor(0x2E, 0x6B, 0x4C)
WARN = RGBColor(0x9C, 0x3E, 0x2E)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
for i, size in [(1, 20), (2, 15), (3, 12.5)]:
    h = doc.styles[f"Heading {i}"]
    h.font.size = Pt(size)
    h.font.color.rgb = RGBColor(0x1C, 0x1B, 0x17)
    h.font.bold = True


def shade(cell, hexc):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hexc)
    cell._tc.get_or_add_tcPr().append(shd)


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 5"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, htxt in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = htxt
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
        shade(c, "EFE2C6")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    doc.add_paragraph()


def bullets(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def numbered(items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def callout(title, body, kind="info"):
    color = {"info": ACCENT, "good": POSITIVE, "warn": WARN}[kind]
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22)
    r = p.add_run(title.upper())
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = color
    b = doc.add_paragraph(body)
    b.paragraph_format.left_indent = Inches(0.22)
    b.paragraph_format.space_after = Pt(12)


def kv(label, text, color=ACCENT):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.1)
    r = p.add_run(label + ": ")
    r.bold = True
    r.font.color.rgb = color
    p.add_run(text)


def usecase(num, name, tag, tagcolor, what, isolate, finding, gain, trade,
            best_for, caveat=None):
    """One full use-case block."""
    doc.add_heading(f"{num}. {name}", level=2)
    p = doc.add_paragraph()
    r = p.add_run(tag)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = tagcolor
    doc.add_paragraph(what)
    kv("Isolate it on the dashboard", isolate)
    kv("Historical finding", finding, POSITIVE)
    kv("Expected gain", gain, POSITIVE)
    kv("How to trade it", trade)
    kv("Best for", best_for)
    if caveat:
        kv("Caveat", caveat, WARN)
    doc.add_paragraph()


# ===========================================================================
# Title
# ===========================================================================
doc.add_heading("Odin's Watchlist — Use Cases & Trading Playbook", level=0)
sub = doc.add_paragraph()
sr = sub.add_run("Every tradeable pattern the tool isolates — what history says about it, "
                 "the gain to expect, and exactly how to trade it")
sr.italic = True
sr.font.size = Pt(13)
doc.add_paragraph()
intro = doc.add_paragraph()
intro.add_run(
    "This is the companion to the main Guide. The Guide explains WHY the tool works; this "
    "document is the field manual: a set of concrete USE CASES you can run day to day. Each "
    "one is a distinct pattern the dashboard can isolate with a filter, backed by what the "
    "backtest actually found, the gain you can reasonably expect, and a step-by-step way to "
    "trade it. They're ordered from highest-confidence to most speculative."
)

callout(
    "How to read the numbers (please read once)",
    "Two different measures appear throughout. HIT RATE = the share of trades that reached the "
    "profit target. BASE RATE = the hit rate of an average watchlist stock (~13% at a +30% "
    "target, ~20% at +20%) — the bar to beat. LIFT = how many times better than random the "
    "top picks were. EXPECTED GAIN (per-trade expectancy) = the average result across many "
    "trades counting winners AND losers together — the number that actually decides if a "
    "pattern makes money. All figures are from one ~13-month window on daily closing prices, "
    "before brokerage/slippage, and ~15% of top picks are volatile BE/BZ names that flatter "
    "the averages. Trust the DIRECTION; treat exact percentages as approximate.",
    kind="warn",
)

callout(
    "The exit rules are shared by EVERY use case below",
    "The exit study was unambiguous: the profit TARGET is a weak lever (anything 25–40% "
    "performs similarly; small targets like +5% actually LOSE money), while the STOP-LOSS and "
    "holding period are the real drivers. So every playbook here uses the same exit spine — a "
    "WIDE target (~25–40%, sized to the stock), a TIGHT stop (~5–6%, widened only for genuinely "
    "jumpy stocks), and a ~2-week time stop on trades that stall. Cutting losers fast is what "
    "turns the edge into profit.",
    kind="good",
)
doc.add_page_break()

# ===========================================================================
# The shared exit engine
# ===========================================================================
doc.add_heading("The shared exit engine (used by all use cases)", level=1)
doc.add_paragraph(
    "Before the individual patterns, here is the exit framework they all plug into. The tool "
    "sizes each stock's stop to its own daily volatility (ATR) and sets the target at ~5x the "
    "risk. The exit study measured the expectancy of each combination on the daily top picks:"
)
table(
    ["Exit setup (on +30% target)", "How often target hit", "Expected gain per trade", "Profit factor"],
    [
        ["+5% target / -8% stop (tempting, DON'T)", "49%", "-0.99% (loses money)", "< 1"],
        ["+20% target / -8% stop (original)", "20%", "+0.26%", "~1.05"],
        ["+30% target / -8% stop", "13%", "+0.80%", "1.17"],
        ["+30% target / -6% stop", "13%", "+1.34%", "1.34"],
        ["+30% target / -5% stop / exit ~2 weeks (best)", "13%", "+1.49%", "1.52"],
    ],
)
callout(
    "The one rule that matters most",
    "A high hit rate is a trap — the +5% target 'wins' 49% of the time and still loses money, "
    "because you're risking 8 to make 5. Judge every pattern by EXPECTED GAIN per trade, not "
    "by how often it wins. And always set the stop BEFORE you buy.",
    kind="warn",
)
doc.add_paragraph(
    "How the per-stock levels come out in practice (target/stop scale with the stock's ATR):"
)
table(
    ["Stock's daily swing (ATR)", "Stop-loss", "Target", "Reward : Risk"],
    [
        ["Calm (<=3.5%)", "-5%", "+25%", "5 : 1"],
        ["Medium (~5%)", "-7.5%", "+38%", "5 : 1"],
        ["Jumpy (~8%)", "-10%", "+40%", "4 : 1"],
        ["Data too fresh to measure", "-6%", "+30%", "5 : 1 (fallback)"],
    ],
)
doc.add_page_break()

# ===========================================================================
# Use cases
# ===========================================================================
doc.add_heading("The use cases", level=1)
doc.add_paragraph(
    "Each block below is self-contained. The dashboard filter named in 'Isolate it' produces "
    "the shortlist; the rest tells you what to expect and how to act."
)

usecase(
    1, "The core shortlist — top ~10 by score", "HIGHEST CONFIDENCE · your daily staple",
    POSITIVE,
    "The everyday use case: take the day's ten highest Setup Scores. The score is led by "
    "relative strength (how much the stock has beaten the market over 20 days), which the "
    "backtest found to be by far the most reliable predictor of a 20%+ move. Its chart-pattern "
    "component is now DATA-DRIVEN too — each screener is weighted by its measured hit-rate and "
    "the dead screeners are zeroed out. The edge lives in this top slice, not in score decimals "
    "— a #3 and a #7 are effectively equal.",
    "Sort by Setup Score, take the top ~10. Keep tiers A and B; ignore C.",
    "On dates the tool was never tuned on (the honest out-of-sample test), the top-10 picks "
    "won 2.2–4.6x more often than random — the edge that survived every honesty check. The "
    "original filter-only score managed only ~1.1–1.4x out-of-sample, so relative strength is "
    "what does the work.",
    "With the shared exit (wide target, ~5–6% stop, 2-week time stop), expectancy is roughly "
    "+1.3% to +1.5% per trade at a +30% target, hit rate ~13% (vs ~13% base — the lift comes "
    "from the top slice being far above the average stock's odds). Most single picks WON'T hit "
    "target; the edge is in playing many with a tight stop.",
    "Buy the ones that also confirm on the chart (real base, breaking out, not extended). Set "
    "the printed per-stock stop before entry. Spread across several names — a string of small "
    "stop-outs is the normal path to the few big winners.",
    "Every trading day — this is the base you start from before applying any lens below.",
    "About 15% of these are volatile BE/BZ (Trade-to-Trade) names whose backtested numbers "
    "flatter the average; scrutinise those extra hard and size them smaller.",
)

usecase(
    2, "Persistent leaders (durable trend)", "HIGH CONFIDENCE · best single filter",
    POSITIVE,
    "Among the strong names, prefer the ones that have STAYED strong. A 'persistent leader' has "
    "led the market (RS >= +15) on 12 or more of the last 20 trading days — a durable trend, "
    "not a one-day flare. You read it straight off the 20-day leadership strip on the dashboard.",
    "Click the 'Persistent leaders' cohort filter (or eyeball a near-full 20-day strip).",
    "Persistent leaders hit target ~24% of the time versus ~15% for freshly-emerged leaders — "
    "the single biggest reliable upgrade among the lenses, and it measures DURABILITY, which "
    "the score alone doesn't.",
    "Raising the hit rate from ~15% to ~24% at a wide target lifts per-trade expectancy "
    "meaningfully above the core — the strongest positive-expectancy cohort the tool offers. "
    "Directionally the best risk-adjusted use case here.",
    "Same exit engine. Prefer a persistent leader that is ALSO firing a fresh breakout today "
    "(see use case 3). If it's resting (Wyckoff BASE), watchlist it rather than chase.",
    "Your core-quality filter — when you want fewer, higher-conviction names from the top slice.",
    "A full strip tells you it's proven, not that today is the entry — still needs a fresh "
    "trigger and a chart check.",
)

usecase(
    3, "Persistent leader + fresh breakout (SOS)", "HIGHEST CONVICTION COMBO",
    POSITIVE,
    "The premium setup: a persistent leader (full strip) that ALSO fires a fresh Wyckoff Sign "
    "of Strength — a breakout out of its range on volume — TODAY. Durability plus a live "
    "trigger on the same name.",
    "Persistent-leaders filter + Wyckoff SOS preset; look for a full strip and an SOS badge "
    "printed today.",
    "This is the ~24% cohort (persistent leaders) intersected with a fresh actionable breakout "
    "— the highest-conviction combination the dashboard can show. Both ingredients "
    "independently beat the base rate; together they concentrate the edge.",
    "The best expected gain of any single-name setup here — durability keeps the hit rate high "
    "while the fresh breakout gives a clean, timeable entry with a tight stop just under the "
    "breakout level.",
    "Enter on the fresh breakout. Stop just below the breakout / base high (usually the printed "
    "~5–6% level). Wide target, let it run, 2-week time stop if it stalls. This can carry a "
    "slightly larger position within your risk rules.",
    "When you want your one or two best ideas of the day.",
    "If it's already up >15% on the day, you're late — wait for a tight pullback or size down "
    "(the 'ext' chase flag).",
)

usecase(
    4, "Episodic Pivot (EP) — the catalyst play", "MEDIUM-HIGH · different axis",
    ACCENT,
    "A stock that GAPS up on heavy volume — the fingerprint of a catalyst (earnings, news, an "
    "order win). It answers a different question from strength: 'did something happen TODAY', "
    "not 'is it already strong'. That independence is what makes it additive.",
    "Click the Episodic Pivot (EP) preset; the EP badge now requires the gap to have HELD "
    "(gap >= 4% on >= 3x volume AND close in the top half of the day's range).",
    "Standalone, EP fires at ~2.5x the base rate. Within the top names it lifts the hit rate to "
    "~24% vs ~17%. A later backtest sharpened it: a gap that HELD (closed top-half of range) "
    "hit target ~19% vs ~12% for a gap that faded (+7pp), so the badge now fires only on held "
    "gaps. Because it's a fresh catalyst rather than a strength proxy, it adds information the "
    "score doesn't already contain.",
    "Comparable expected gain to the persistent-leader cohort when it appears within the top "
    "names — a genuinely positive-expectancy pattern, with the bonus of being early.",
    "Enter on the gap day if the gap comes off a tight base and HOLDS above the gap level. Stop "
    "below the gap/base. Catalyst moves can continue hard or fade fast, so the stop matters "
    "more than usual; keep the target wide.",
    "Days with real news flow; pairing a catalyst with existing strength.",
    "A gap on a stock with a 0/20 strip and no base is a speculative one-off — size it like "
    "use case 7, not like a leader.",
)

usecase(
    5, "Wyckoff SOS breakout (structure entry)", "MEDIUM · clean timing",
    ACCENT,
    "Trade the breakout itself: a Wyckoff Sign of Strength is a stock 'jumping the creek' — "
    "breaking out of an accumulation range on expanding volume. It's the classic actionable "
    "breakout entry, inferred from which screeners the stock fired.",
    "Wyckoff SOS preset (fresh breakouts). Combine with a full leadership strip for quality.",
    "Wyckoff breakout survived the backtest as a lens that measures STRUCTURE (a different axis "
    "from raw strength) — unlike the Minervini trend template and delivery-% accumulation, "
    "which were tested and REJECTED for merely re-measuring strength. Strongest when the "
    "breakout stock is also a leader (folds into use cases 2–3).",
    "Positive expectancy when combined with strength/leadership; on its own it's a timing tool "
    "that gives a clean, tight-stop entry rather than a standalone edge.",
    "Buy the breakout, stop just under the range high / breakout level, wide target, 2-week "
    "time stop. Skip breakouts with no leadership history behind them.",
    "Timing an entry on a name you already like from the score/strip.",
    "A breakout with a 0/20 strip is a flare, not a trend — confirm leadership first.",
)

usecase(
    6, "Wyckoff LPS pullback (lower-chase entry)", "MEDIUM · patient entry",
    ACCENT,
    "The Last Point of Support: a quiet, low-volume pullback to support AFTER a breakout. It's "
    "the safer, lower-chase Wyckoff entry — you buy the retest instead of the breakout candle.",
    "Wyckoff LPS preset; look for a post-breakout stock easing back on light volume.",
    "The same structural family as SOS and validated on the same basis (a different axis from "
    "strength). Its value is a better ENTRY PRICE and tighter stop on an already-strong name, "
    "reducing the chase risk that hurts breakout buyers.",
    "Improves expectancy mainly by lowering entry cost and tightening the stop on a leader you "
    "were going to trade anyway — not a separate edge so much as a cheaper way into use cases "
    "2–3.",
    "Buy the retest as it holds support on shrinking volume; stop just below support. Same wide "
    "target and time stop. Ideal when you missed the initial SOS breakout.",
    "Entering a strong name without paying up for a hot breakout.",
    "If the 'pullback' comes on HEAVY volume, it may be a real breakdown — stand aside.",
)

usecase(
    7, "Emerging (early move) — catch the wake-up", "SPECULATIVE · early, lower-confidence",
    WARN,
    "Catch a stock a step BEFORE it's an obvious leader: a volume THRUST (>= 3x its own normal) "
    "while its relative strength is rising but still modest (0 to +25). The moment a quiet "
    "stock wakes up — the STALLION-type early move.",
    "Click the 'Emerging' (EMRG) cohort filter.",
    "Tested honestly: buying BEFORE strength shows up (the quiet pullback) is NO better than "
    "random — you'd be front-running the only signal that works. But catching the THRUST itself "
    "with RS accelerating hits target at ~1.5x the base rate. Real, but roughly HALF the edge "
    "of a confirmed leader (~3x).",
    "Positive but modest expectancy — about half the per-trade edge of use cases 1–3. You're "
    "paying for earliness with lower reliability.",
    "Treat it as speculative: SMALLER position, only if the chart shows the thrust coming off a "
    "tight base and holding. Same tight stop and wide target. Let the confirmed-leader cohorts "
    "remain your core.",
    "A satellite position when you want early exposure and accept lower odds.",
    "Never make this your core allocation — it is explicitly a lower-confidence entry, not a "
    "better one.",
)

usecase(
    8, "RIGHTWAY Telegram confluence", "CONFIRMING NUDGE · not yet backtested",
    WARN,
    "Cross-reference the dashboard's picks against your RIGHTWAY Telegram group. When the tool "
    "independently flags a strong setup AND your group is already on it, that's a second, "
    "unrelated vote — pure confluence, never a signal by itself.",
    "'In RIGHTWAY' filter (anyone mentioned it) or the tighter 'Admin call' filter (the "
    "broadcast channel actually called it, shown by a '^' and a '*' for recent conviction "
    "tags). Run telegram_fetch.py then parse_telegram_recs.py at end of day first.",
    "NOT yet backtested — a proper test of whether the admin's calls PRECEDE moves (rather than "
    "trail them) is the planned next step. What we DO know: the channel is promotional and "
    "often posts winners after they've moved, and ~43 of 60 dashboard names get mentioned on a "
    "given day, so a plain mention isn't selective.",
    "No measured expectancy yet — do not assign one. Value is qualitative confluence, best when "
    "it's 'strong tool setup + FRESH admin call', not raw mention count.",
    "Use only to break ties or add conviction to a stock the tool ALREADY rates. Never buy "
    "something the tool doesn't like just because the group is talking about it. The tool's own "
    "volatility-based target/stop are the ones with a backtest — the admin's target is their "
    "goal, and most admin calls in the data are 9–24 months old (historical context, not live).",
    "A final confirming glance after use cases 1–3 have already selected a name.",
    "Promotional source. Treat every mention as 'they're discussing it', not 'it's a buy'.",
)

# ===========================================================================
# What NOT to trade
# ===========================================================================
doc.add_heading("Anti-patterns — use cases we tested and REJECTED", level=1)
doc.add_paragraph(
    "Part of trusting the good use cases is seeing what failed the same test. These were built, "
    "backtested, and deliberately dropped — recognising them saves you from bad trades:"
)
table(
    ["Rejected pattern", "Why it failed the backtest / what to do instead"],
    [
        ["High score but 0/20 leadership strip",
         "A one-day flare with no durable trend behind it (e.g. EMAMIPAP: score 92, RS +38, "
         "yet never a leader in 20 days). The high score is necessary, not sufficient — the "
         "empty strip is the tell. Skip it."],
        ["Already up >15% today ('ext' chase flag)",
         "You'd be buying the spike late; the risk:reward from here is poor. Wait for a tight "
         "pullback or pass."],
        ["Small profit targets (+5% to +10%)",
         "They 'win' often but LOSE money — risking an 8% stop to make 5% is negative "
         "expectancy. Always use a wide target."],
        ["Minervini trend template (price above long-term averages)",
         "Tested and dropped — it merely re-measures what relative strength already captures. "
         "Redundant, not additive."],
        ["NSE delivery-% accumulation",
         "Tested and dropped — high-delivery names actually continued slightly WORSE. Another "
         "strength proxy, no independent edge."],
        ["Buying the quiet pre-thrust dip (front-running Emerging)",
         "No better than random — you're guessing ahead of the only signal that works. Wait "
         "for the actual volume thrust."],
    ],
)
callout(
    "The pattern behind every accepted vs rejected use case",
    "Everything that SURVIVED (persistence, Wyckoff structure, episodic pivot) measures a "
    "DIFFERENT axis — durability, structure, or a catalyst. Everything REJECTED (Minervini, "
    "delivery %) was just another proxy for strength the score already has. When you consider a "
    "new idea, ask: does it add a new axis, or re-measure strength? Only the former helps.",
    kind="good",
)

# ===========================================================================
# Daily routine
# ===========================================================================
doc.add_heading("Putting it together — your daily routine", level=1)
numbered([
    "After the close, run ONE command: run_odin.py. It now chains the whole flow automatically "
    "— pulls new RIGHTWAY Telegram messages, scores the sheet, updates the paper ledger, and "
    "rebuilds the dashboard (with its P&L panel).",
    "Open the dashboard and start from the top ~10 by score — use case 1. Glance at the CHANGES "
    "tab to see what's new or has moved since yesterday, and set your capital + risk-% in the "
    "top bar so every row shows the share size to buy.",
    "Filter to your conviction: Persistent leaders (2), then look for a fresh SOS on them (3). "
    "Scan EP (4) for catalysts — remember the badge now only shows gaps that HELD. These are "
    "your primary trades. Star the ones you like (they collect in the WATCHLIST tab) and click "
    "'TV' to confirm the chart.",
    "Optionally add ONE speculative Emerging name (7) at smaller size.",
    "Glance at RIGHTWAY confluence (8) only to confirm names you already like.",
    "Check the dashboard's P&L panel to see how recent picks — and each lens — are actually "
    "doing (the paper ledger). Let it accumulate a few weeks before reading much into it.",
    "Reject anything with a 0/20 strip or an 'ext' flag, however high the score.",
    "For each keeper: confirm on the chart, set the printed stop BEFORE buying, use the wide "
    "target, and honour the ~2-week time stop on stalled trades.",
    "Size so a run of small stop-outs is survivable — that string of small losers punctuated "
    "by a few big winners IS the strategy working, not failing.",
])

callout(
    "The whole playbook in three sentences",
    "The SCORE (relative strength) gets a name onto the list; the LENSES (persistence, "
    "Wyckoff, EP, Emerging, RIGHTWAY) tell you which of the strong names is a buy, a wait, or a "
    "trap; the CHART is the final confirmation. Expected gain is positive but small per trade "
    "and comes from many trades with a tight stop — never from any single pick. Cut losers at "
    "the stop, let winners run to a wide target: that discipline is the edge.",
    kind="good",
)

doc.add_paragraph()
foot = doc.add_paragraph()
fr = foot.add_run(
    "Companion to Odins_Watchlist_Guide.docx. All figures from the project's own backtests: "
    "2,554 watchlist stocks over 158 trading days (Jul 2025 – Jul 2026), price history over "
    "528 unbroken days (Jul 2024 – Jul 2026). Directional, not guarantees; daily closes, before "
    "costs; ~15% of picks are volatile BE/BZ names. Always use a stop-loss."
)
fr.italic = True
fr.font.size = Pt(9)
fr.font.color.rgb = INK_SOFT

OUT = "/Users/vinaykusuma/Documents/AlgoTrade/odins_watchlist/working_version/Odins_Watchlist_UseCases.docx"
doc.save(OUT)
print("Saved:", OUT)
